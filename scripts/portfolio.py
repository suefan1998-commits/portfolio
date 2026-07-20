#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import html
import json
import os
import re
import shlex
import shutil
import subprocess
import sys
import textwrap
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

try:
    import openpyxl
except Exception as exc:  # pragma: no cover
    raise SystemExit("缺少 openpyxl，无法读取 portfolio_registry.xlsx。") from exc


ROOT = Path(__file__).resolve().parents[1]
REGISTRY = ROOT / "portfolio_registry.xlsx"
INPUT = ROOT / "input"
OUTPUT = ROOT / "output"
WORK = ROOT / "work"
INBOX = INPUT / "inbox"
CONTENT = OUTPUT / "content"
SUMMARIES = OUTPUT / "summaries"
PUBLIC = OUTPUT / "site"
ASSETS = INPUT / "assets"
PDF_DIR = OUTPUT / "pdf"
TMP = WORK / "tmp"
TENCENT_ENV = ROOT / ".env.tencent-server"
NODE = Path("/Users/Sue/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node")
NODE_MODULES = Path("/Users/Sue/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/node_modules")
PNPM = Path("/Users/Sue/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/fallback/pnpm")
WORKBOOK_TOOLS = ROOT / "scripts" / "workbook_tools.mjs"

REQUIRED_HEADERS = [
    "作品ID",
    "标题",
    "Keyline",
    "来源类型",
    "链接/文件名",
    "撰稿类型",
    "主题标签",
    "发布媒体",
    "发布日期",
    "是否公开展示",
    "是否为精选作品",
    "处理状态",
    "备注",
]
OPTIONAL_HEADERS = [
    "作品图片",
]
HEADERS = REQUIRED_HEADERS + OPTIONAL_HEADERS

MIN_WEB_TEXT_LENGTH = 180
PDF_FONT = "PortfolioCJK"
PDF_FONT_BOLD = "PortfolioCJKBold"


@dataclass
class RegistryRow:
    row_index: int
    work_id: str
    title: str
    keyline: str
    source_type: str
    source: str
    writing_type: str
    tags: str
    media: str
    publish_date: str
    public: str
    featured: str
    status: str
    notes: str
    article_image: str


class ArticleHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.skip_depth = 0
        self.current_tag = ""
        self.title_chunks: list[str] = []
        self.text_chunks: list[str] = []
        self.line_parts: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:
        self.current_tag = tag.lower()
        if self.current_tag in {"script", "style", "noscript", "svg", "canvas"}:
            self.skip_depth += 1
        if self.current_tag == "br":
            self.flush_line()
        elif self.current_tag in {"p", "h1", "h2", "h3", "blockquote", "li", "article", "section", "div"}:
            self.flush_line()

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg", "canvas"} and self.skip_depth:
            self.skip_depth -= 1
        if self.current_tag == tag:
            self.current_tag = ""
        if tag in {"p", "h1", "h2", "h3", "blockquote", "li", "article", "section", "div"}:
            self.flush_line()

    def handle_data(self, data: str) -> None:
        if self.skip_depth:
            return
        text = normalize_spaces(data)
        if not text:
            return
        if self.current_tag == "title":
            self.title_chunks.append(text)
        else:
            self.line_parts.append(text)

    def close(self) -> None:
        self.flush_line()
        super().close()

    def flush_line(self) -> None:
        line = normalize_spaces(" ".join(self.line_parts))
        self.line_parts = []
        if line:
            self.text_chunks.append(line)


def normalize_spaces(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\u3000", " ")
    return re.sub(r"\s+", " ", text).strip()


def normalize_date(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, dt.datetime | dt.date):
        return value.strftime("%Y-%m-%d")
    text = normalize_spaces(value)
    if not text:
        return ""
    match = re.search(r"(\d{4})[./-](\d{1,2})[./-](\d{1,2})", text)
    if match:
        year, month, day = match.groups()
        return f"{int(year):04d}-{int(month):02d}-{int(day):02d}"
    return text


def split_tags(text: str) -> list[str]:
    if re.search(r"#[^\s#]+(?:\s+#[^\s#]+)+", text or ""):
        parts = re.findall(r"#[^\s#]+", text)
    else:
        parts = re.split(r"[、,，;；/｜|\s]+", text or "")
    return [part.strip() for part in parts if part.strip()]


def slugify(text: str, fallback: str = "item") -> str:
    text = normalize_spaces(text).lower()
    ascii_slug = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    if ascii_slug:
        return ascii_slug[:80]
    code = "-".join(f"{ord(ch):x}" for ch in text[:12] if ch.strip())
    return code or fallback


def ensure_dirs() -> None:
    for folder in [INBOX, CONTENT, SUMMARIES, PUBLIC, ASSETS, PDF_DIR, TMP]:
        folder.mkdir(parents=True, exist_ok=True)


def ensure_node_modules() -> None:
    link = ROOT / "node_modules"
    if link.exists() or link.is_symlink():
        return
    if not NODE_MODULES.exists():
        raise SystemExit("找不到工作簿运行依赖，无法创建或更新 portfolio_registry.xlsx。")
    link.symlink_to(NODE_MODULES, target_is_directory=True)


def run_node(*args: str) -> None:
    ensure_node_modules()
    node = NODE if NODE.exists() else shutil.which("node")
    if not node:
        raise SystemExit("找不到 Node.js，无法运行工作簿辅助工具。")
    subprocess.run([str(node), str(WORKBOOK_TOOLS), *args], cwd=ROOT, check=True)


def init_project(force: bool = False) -> None:
    ensure_dirs()
    if REGISTRY.exists() and not force:
        print(f"台账已存在：{REGISTRY}")
        return
    run_node("create-template", str(REGISTRY))
    print(f"已生成台账模板：{REGISTRY}")


def load_registry() -> list[RegistryRow]:
    if not REGISTRY.exists():
        raise SystemExit("未找到 portfolio_registry.xlsx。请先运行：python3 scripts/portfolio.py init")
    workbook = openpyxl.load_workbook(REGISTRY, data_only=True)
    if "作品台账" not in workbook.sheetnames:
        raise SystemExit("portfolio_registry.xlsx 缺少工作表：作品台账")
    sheet = workbook["作品台账"]
    headers = [normalize_spaces(cell.value) for cell in sheet[1]]
    missing = [header for header in REQUIRED_HEADERS if header not in headers]
    if missing:
        raise SystemExit("台账缺少字段：" + "、".join(missing))

    col = {name: headers.index(name) + 1 for name in headers if name}
    def cell_value(index: int, name: str) -> str:
        if name not in col:
            return ""
        return normalize_spaces(sheet.cell(index, col[name]).value)

    rows: list[RegistryRow] = []
    for index in range(2, sheet.max_row + 1):
        row_values = [sheet.cell(index, col[name]).value for name in REQUIRED_HEADERS]
        if not any(normalize_spaces(value) for value in row_values):
            continue
        rows.append(
            RegistryRow(
                row_index=index,
                work_id=cell_value(index, "作品ID"),
                title=cell_value(index, "标题"),
                keyline=cell_value(index, "Keyline"),
                source_type=cell_value(index, "来源类型"),
                source=cell_value(index, "链接/文件名"),
                writing_type=cell_value(index, "撰稿类型"),
                tags=cell_value(index, "主题标签"),
                media=cell_value(index, "发布媒体"),
                publish_date=normalize_date(sheet.cell(index, col["发布日期"]).value),
                public=cell_value(index, "是否公开展示"),
                featured=cell_value(index, "是否为精选作品"),
                status=cell_value(index, "处理状态"),
                notes=cell_value(index, "备注"),
                article_image=cell_value(index, "作品图片"),
            )
        )
    return rows


def make_work_id(row: RegistryRow, title: str = "") -> str:
    if row.work_id and not is_legacy_work_id(row.work_id):
        return row.work_id
    prefix = row.publish_date.replace("-", "") if row.publish_date else dt.date.today().strftime("%Y%m%d")
    media = sanitize_filename_part(row.media) or "未标注媒体"
    title_part = sanitize_filename_part(title or row.title or row.source) or "未命名作品"
    return f"{prefix}-{media}-{title_part}"[:140]


def is_legacy_work_id(work_id: str) -> bool:
    parts = work_id.split("-")
    return len(parts) > 3 and all(re.fullmatch(r"[0-9a-f]+", part) for part in parts[1:4])


def sanitize_filename_part(text: object) -> str:
    text = normalize_spaces(text)
    text = re.sub(r"[^\w\u4e00-\u9fff]+", "", text, flags=re.UNICODE)
    return text.replace("_", "")


def fetch_web_article(url: str) -> tuple[str, str]:
    request = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 PortfolioArchiver/1.0",
            "Accept": "text/html,application/xhtml+xml",
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=20) as response:
            charset = response.headers.get_content_charset() or "utf-8"
            raw = response.read(5_000_000)
    except urllib.error.URLError as exc:
        raise RuntimeError(f"链接无法访问：{exc}") from exc

    html_text = raw.decode(charset, errors="replace")
    content_html = extract_preferred_article_html(html_text) or html_text
    parser = ArticleHTMLParser()
    parser.feed(content_html)
    parser.close()

    title = extract_html_title(html_text) or " ".join(parser.title_chunks)
    title = re.sub(r"[_|-].*$", "", title).strip()
    chunks = dedupe_chunks(parser.text_chunks)
    text = "\n\n".join(chunks)
    return title, html.unescape(text)


def extract_html_title(html_text: str) -> str:
    patterns = [
        r'<meta\s+property=["\']og:title["\']\s+content=["\']([^"\']+)["\']',
        r'<meta\s+content=["\']([^"\']+)["\']\s+property=["\']og:title["\']',
        r'var\s+msg_title\s*=\s*["\'](.*?)["\']',
        r'<title[^>]*>(.*?)</title>',
    ]
    for pattern in patterns:
        match = re.search(pattern, html_text, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            continue
        value = decode_embedded_string(match.group(1))
        value = re.sub(r"<[^>]+>", "", value)
        value = normalize_spaces(html.unescape(value))
        if value:
            return value
    return ""


def decode_embedded_string(value: str) -> str:
    try:
        return json.loads(f'"{value}"')
    except Exception:
        return value.replace(r"\/", "/")


def extract_preferred_article_html(html_text: str) -> str:
    markers = [
        'id="js_content"',
        "id='js_content'",
        "rich_media_content",
        "js_underline_content",
    ]
    positions = [html_text.find(marker) for marker in markers if html_text.find(marker) >= 0]
    if not positions:
        return ""
    marker_index = min(positions)
    start = html_text.rfind("<", 0, marker_index)
    if start < 0:
        return ""
    return extract_balanced_element(html_text, start)


def extract_balanced_element(html_text: str, start: int) -> str:
    first = re.match(r"<\s*([a-zA-Z0-9]+)\b", html_text[start:])
    if not first:
        return ""
    root_tag = first.group(1).lower()
    void_tags = {"area", "base", "br", "col", "embed", "hr", "img", "input", "link", "meta", "source", "track", "wbr"}
    tag_pattern = re.compile(r"<\s*(/)?\s*([a-zA-Z0-9]+)\b[^>]*?>", re.DOTALL)
    depth = 0
    for match in tag_pattern.finditer(html_text, start):
        closing, tag = match.group(1), match.group(2).lower()
        if tag != root_tag:
            continue
        raw_tag = match.group(0)
        if closing:
            depth -= 1
            if depth == 0:
                return html_text[start : match.end()]
        elif tag not in void_tags and not raw_tag.rstrip().endswith("/>"):
            depth += 1
    return html_text[start:]


def dedupe_chunks(chunks: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    credits: list[str] = []
    credit_seen: set[str] = set()
    for chunk in chunks:
        clean = normalize_spaces(html.unescape(chunk))
        if len(clean) < 8:
            continue
        if is_editorial_credit(clean):
            if clean not in credit_seen:
                credit_seen.add(clean)
                credits.append(clean)
            continue
        if is_caption_like(clean) and not is_editorial_credit(clean):
            continue
        key = clean[:120]
        if key in seen:
            continue
        seen.add(key)
        result.append(clean)
    result.extend(credits)
    return result


def is_caption_like(text: str) -> bool:
    clean = normalize_spaces(text)
    if not clean:
        return True
    if len(clean) > 80:
        return False

    credit_patterns = [
        r"均为\s*[A-Za-z]",
        r"(图源|图片|供图|摄影|Photo|Courtesy|©|来源)[:：]?",
        r"(外套|针织衫|长裤|衬衫|夹克|裙|鞋|腕表|包)\s*[、，,]",
        r"(滑动|查看更多|原文刊登|采访\s*&|部分图片)",
        r"\d+(\.\d+)?\s*cm\s*[x×]\s*\d+",
    ]
    if any(re.search(pattern, clean, re.IGNORECASE) for pattern in credit_patterns):
        return True

    artwork_mediums = (
        "琉璃|玻璃|水晶|铜|铁|木|石|陶|瓷|纸本|布面|油画|"
        "雕塑|装置|影像|摄影|建筑|展厅|美术馆|材料"
    )
    if "《" in clean and "》" in clean and re.search(rf"[，,、]\s*({artwork_mediums})", clean):
        return True

    scene_caption_words = "日落|风貌|外观|内景|展厅|现场|远景|近景|合影|剧照|海报"
    sentence_end = re.search(r"[。！？!?；;]$", clean)
    if len(clean) <= 36 and not sentence_end and re.search(scene_caption_words, clean):
        return True

    has_content_marker = re.search(r"[。！？!?：:；;]|\?|\？", clean)
    if len(clean) <= 42 and not has_content_marker:
        return True

    return False


def is_editorial_credit(text: str) -> bool:
    clean = normalize_spaces(text)
    if len(clean) > 120:
        return False
    return bool(re.search(r"(撰稿|撰文|作者|编辑|责编)\s*[:：／/]", clean))


def extract_docx(path: Path) -> str:
    document = Document(path)
    pieces: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_spaces(paragraph.text)
        if text:
            pieces.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [normalize_spaces(cell.text) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                pieces.append(line)
    return "\n\n".join(pieces)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_text_file(path: Path) -> str:
    for encoding in ["utf-8", "utf-8-sig", "gb18030"]:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def extract_print_article(filename: str) -> str:
    path = INBOX / filename
    if not path.exists():
        raise RuntimeError(f"input/inbox 中找不到文件：{filename}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md"}:
        return extract_text_file(path)
    raise RuntimeError(f"暂不支持的纸刊文件格式：{suffix}")


def frontmatter(row: RegistryRow, work_id: str, title: str, source_kind: str) -> str:
    data = {
        "id": work_id,
        "title": title,
        "keyline": row.keyline,
        "source_type": row.source_type,
        "source": row.source,
        "writing_type": row.writing_type,
        "tags": split_tags(row.tags),
        "media": row.media,
        "publish_date": row.publish_date,
        "public": row.public or "是",
        "featured": row.featured or "否",
        "article_image": row.article_image,
        "source_kind": source_kind,
        "notes": row.notes,
    }
    lines = ["---"]
    for key, value in data.items():
        if isinstance(value, list):
            lines.append(f"{key}: {json.dumps(value, ensure_ascii=False)}")
        else:
            escaped = str(value).replace('"', '\\"')
            lines.append(f'{key}: "{escaped}"')
    lines.append("---")
    return "\n".join(lines)


def write_article(row: RegistryRow, title: str, text: str, source_kind: str) -> str:
    work_id = make_work_id(row, title)
    output = CONTENT / f"{work_id}.md"
    output.write_text(
        frontmatter(row, work_id, title, source_kind) + "\n\n" + text.strip() + "\n",
        encoding="utf-8",
    )
    return work_id


def summary_path(article: dict) -> Path:
    return SUMMARIES / f"{article.get('id')}.md"


def summary_frontmatter(article: dict, status: str = "待审核") -> str:
    data = {
        "id": article.get("id", ""),
        "title": article.get("title", ""),
        "source_type": article.get("source_type", ""),
        "source": article.get("source", ""),
        "writing_type": article.get("writing_type", ""),
        "tags": article.get("tags", []),
        "media": article.get("media", ""),
        "publish_date": article.get("publish_date", ""),
        "public": article.get("public", "是"),
        "featured": article.get("featured", "否"),
        "article_image": article.get("article_image", ""),
        "source_kind": article.get("source_kind", ""),
        "summary_status": status,
        "keyline": article.get("keyline", ""),
        "notes": article.get("notes", ""),
    }
    lines = ["---"]
    for key, value in data.items():
        if isinstance(value, list):
            lines.append(f"{key}: {json.dumps(value, ensure_ascii=False)}")
        else:
            escaped = str(value).replace('"', '\\"')
            lines.append(f'{key}: "{escaped}"')
    lines.append("---")
    return "\n".join(lines)


def split_article_sections(body: str) -> list[tuple[str, list[str]]]:
    blocks = [item.strip() for item in body.split("\n\n") if item.strip()]
    has_headings = any(markdown_heading(block) for block in blocks)
    sections: list[tuple[str, list[str]]] = []
    current_title = "" if has_headings else "精选节选"
    current_paragraphs: list[str] = []
    for block in blocks:
        heading = markdown_heading(block)
        if heading:
            if current_paragraphs:
                sections.append((current_title, current_paragraphs))
            current_title = heading[1]
            current_paragraphs = []
            continue
        if is_editorial_credit(block):
            continue
        if len(normalize_spaces(block)) < 45:
            continue
        current_paragraphs.append(block)
    if current_paragraphs:
        sections.append((current_title, current_paragraphs))
    return sections


def paragraph_score(paragraph: str) -> int:
    text = normalize_spaces(paragraph)
    score = 0
    if 90 <= len(text) <= 360:
        score += 3
    if re.search(r"(然而|但|因此|这|正是|意味着|问题|核心|背后|成为|并非|不是|而是)", text):
        score += 2
    if re.search(r"(城市|文化|设计|人物|品牌|社会|旅行|建筑|商业|技术|社区)", text):
        score += 1
    if len(text) > 520:
        score -= 2
    return score


def draft_overview(article: dict, sections: list[tuple[str, list[str]]]) -> str:
    title = str(article.get("title") or "这篇作品")
    first = ""
    for _, paragraphs in sections:
        if paragraphs:
            first = trim_text(paragraphs[0], 58).rstrip("。")
            break
    if first:
        return f"围绕《{title}》，文章呈现{first}。"
    return f"围绕《{title}》，文章梳理其核心议题与文本线索。"


def draft_summary_body(article: dict) -> str:
    sections = split_article_sections(str(article.get("body") or ""))
    lines = []
    for title, paragraphs in sections:
        ranked = sorted(paragraphs, key=paragraph_score, reverse=True)[:2]
        if not ranked:
            continue
        if title:
            lines.extend(["", f"## {title}"])
        for paragraph in ranked:
            lines.extend(["", f"……{trim_text(paragraph, 260).strip('。')}。……"])
    return "\n".join(lines).strip()


def ensure_summary_draft(article: dict) -> bool:
    path = summary_path(article)
    if path.exists():
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    draft = dict(article)
    draft["keyline"] = normalize_spaces(article.get("keyline")) or draft_overview(article, split_article_sections(str(article.get("body") or "")))
    path.write_text(
        summary_frontmatter(draft, "待审核") + "\n\n" + draft_summary_body(article) + "\n",
        encoding="utf-8",
    )
    return True


def generate_missing_summaries() -> tuple[int, int, int]:
    ensure_dirs()
    created = 0
    pending = 0
    confirmed = 0
    for article in load_articles(public_only=False):
        if ensure_summary_draft(article):
            created += 1
        summary = parse_article(summary_path(article)) if summary_path(article).exists() else {}
        status = normalize_spaces(summary.get("summary_status"))
        if status == "已确认":
            confirmed += 1
        else:
            pending += 1
    return created, pending, confirmed


def summaries_status() -> None:
    created, pending, confirmed = generate_missing_summaries()
    print(f"summary 草稿检查完成：新增 {created} 篇，待审核 {pending} 篇，已确认 {confirmed} 篇。")
    for path in sorted(SUMMARIES.glob("*.md")):
        if path.name == ".gitkeep":
            continue
        article = parse_article(path)
        status = normalize_spaces(article.get("summary_status")) or "待审核"
        if status != "已确认":
            print(f"- {status}：{article.get('title') or path.stem}")


def update_registry_status(updates: list[dict]) -> None:
    if not updates:
        return
    TMP.mkdir(exist_ok=True)
    updates_path = TMP / "registry_updates.json"
    updates_path.write_text(json.dumps(updates, ensure_ascii=False, indent=2), encoding="utf-8")
    run_node("update-status", str(REGISTRY), str(updates_path))


def ingest() -> None:
    ensure_dirs()
    rows = load_registry()
    updates: list[dict] = []
    processed = 0
    for row in rows:
        if row.status and row.status not in {"待处理", "需手动补充", "错误"}:
            continue
        if not row.source_type or not row.source:
            updates.append(
                {
                    "rowIndex": row.row_index,
                    "status": "错误",
                    "note": append_note(row.notes, "缺少来源类型或链接/文件名"),
                }
            )
            continue
        if not row.keyline:
            updates.append(
                {
                    "rowIndex": row.row_index,
                    "status": "错误",
                    "note": append_note(row.notes, "缺少 Keyline"),
                }
            )
            continue

        try:
            if row.source_type == "新媒体":
                title, text = fetch_web_article(row.source)
                title = row.title or title or row.source
                if len(text) < MIN_WEB_TEXT_LENGTH:
                    updates.append(
                        {
                            "rowIndex": row.row_index,
                            "title": title,
                            "status": "需手动补充",
                            "note": append_note(row.notes, "网页正文过短或平台限制，需手动补充正文"),
                        }
                    )
                    continue
                source_kind = "web"
            elif row.source_type == "纸刊":
                text = extract_print_article(row.source)
                title = row.title or Path(row.source).stem
                if not text.strip():
                    raise RuntimeError("未提取到正文")
                source_kind = "print"
            else:
                raise RuntimeError("来源类型必须是“新媒体”或“纸刊”")

            work_id = write_article(row, title, text, source_kind)
            updates.append(
                {
                    "rowIndex": row.row_index,
                    "id": work_id,
                    "title": title,
                    "status": "已入库",
                    "note": append_note(row.notes, "已整理入库"),
                }
            )
            processed += 1
        except Exception as exc:
            updates.append(
                {
                    "rowIndex": row.row_index,
                    "status": "错误",
                    "note": append_note(row.notes, str(exc)),
                }
            )

    update_registry_status(updates)
    print(f"入库完成：新增或更新 {processed} 篇，状态更新 {len(updates)} 条。")
    if processed:
        print("下一步：请在对话中确认 AI 生成的 keyline 与章节精选后，再写入 summary。")


def append_note(old_note: str, new_note: str) -> str:
    stamp = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    entry = f"[{stamp}] {new_note}"
    if not old_note:
        return entry
    if new_note in old_note:
        return old_note
    return f"{old_note}\n{entry}"


def parse_article(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        raise ValueError(f"作品文件缺少元信息：{path.name}")
    _, meta_text, body = text.split("---", 2)
    meta: dict[str, object] = {}
    for line in meta_text.strip().splitlines():
        if ":" not in line:
            continue
        key, raw_value = line.split(":", 1)
        raw_value = raw_value.strip()
        if raw_value.startswith("["):
            value = json.loads(raw_value)
        elif raw_value.startswith('"') and raw_value.endswith('"'):
            value = raw_value[1:-1].replace('\\"', '"')
        else:
            value = raw_value
        meta[key.strip()] = value
    meta["body"] = body.strip()
    meta["path"] = path
    return meta


def load_articles(public_only: bool = False) -> list[dict]:
    overrides = registry_article_overrides()
    articles = []
    for path in sorted(CONTENT.glob("*.md")):
        if path.name == ".gitkeep":
            continue
        article = parse_article(path)
        if article.get("id") in overrides:
            for key, value in overrides[article["id"]].items():
                if value != "":
                    article[key] = value
        if public_only and article.get("public", "是") != "是":
            continue
        articles.append(article)
    articles.sort(key=lambda item: str(item.get("publish_date") or ""), reverse=True)
    return articles


def load_summary_articles(public_only: bool = False) -> list[dict]:
    articles = []
    for article in load_articles(public_only=public_only):
        item = dict(article)
        item["_full_body"] = article.get("body", "")
        path = summary_path(article)
        if path.exists():
            summary = parse_article(path)
            summary_body = str(summary.get("body") or "")
            for key, value in summary.items():
                if key not in {"body", "path"} and value != "":
                    item[key] = value
            item["summary_status"] = normalize_spaces(summary.get("summary_status")) or "待审核"
            item["summary_body"] = summary_body
            item["body"] = summary_body if item["summary_status"] == "已确认" else ""
            item["summary_path"] = path
        else:
            item["summary_status"] = "待生成"
            item["summary_body"] = ""
            item["body"] = ""
            item["summary_path"] = path
        articles.append(item)
    articles.sort(key=lambda item: str(item.get("publish_date") or ""), reverse=True)
    return articles


def confirmed_summary(article: dict) -> bool:
    return normalize_spaces(article.get("summary_status")) == "已确认"


def public_body(article: dict) -> str:
    return str(article.get("body") or "") if confirmed_summary(article) else ""


def public_preview(article: dict, limit: int = 220) -> str:
    keyline = normalize_spaces(article.get("keyline"))
    if confirmed_summary(article) and keyline:
        return trim_text(keyline, limit)
    body = public_body(article)
    if not body:
        return "摘要待确认。请点击阅读原文查看完整作品。"
    paragraphs = []
    for block in body.split("\n\n"):
        text = block.strip()
        if not text or re.match(r"^\s{0,3}#{1,6}\s+", text):
            continue
        if text.startswith("【概述】"):
            text = text.removeprefix("【概述】").strip()
        if text:
            paragraphs.append(text)
    return trim_text(paragraphs[0], limit) if paragraphs else "摘要待确认。请点击阅读原文查看完整作品。"


def registry_article_overrides() -> dict[str, dict[str, str]]:
    if not REGISTRY.exists():
        return {}
    overrides: dict[str, dict[str, str]] = {}
    try:
        rows = load_registry()
    except Exception:
        return overrides
    for row in rows:
        if not row.work_id:
            continue
        overrides[row.work_id] = {
            "source": row.source,
            "keyline": row.keyline,
            "article_image": row.article_image,
            "featured": row.featured,
            "public": row.public,
        }
    return overrides


def profile_image_path() -> str:
    path = ASSETS / "profile.jpg"
    return str(path.relative_to(ROOT)) if path.exists() else ""


def html_escape(text: object) -> str:
    return html.escape(str(text or ""), quote=True)


def is_remote_asset(value: str) -> bool:
    return bool(re.match(r"https?://", normalize_spaces(value), flags=re.IGNORECASE))


def original_url(article: dict) -> str:
    source = normalize_spaces(article.get("source"))
    if re.match(r"https?://", source, flags=re.IGNORECASE):
        return source
    return ""


def source_link_html(article: dict, class_name: str = "source-link") -> str:
    url = original_url(article)
    if not url:
        return ""
    safe_url = html_escape(url)
    return f'<p class="{class_name}"><a href="{safe_url}" target="_blank" rel="noopener noreferrer">阅读原文</a></p>'


def article_source_prompt_html(article: dict) -> str:
    url = original_url(article)
    if not url:
        return ""
    safe_url = html_escape(url)
    return (
        '<p class="article-source-link">'
        f'完整内容请点击 <a href="{safe_url}" target="_blank" rel="noopener noreferrer">阅读原文</a>，以下为精选节选'
        "</p>"
    )


def display_date(value: object) -> str:
    return normalize_spaces(value)


def display_date_dot(value: object) -> str:
    return display_date(value).replace("-", ".")


def date_sort_desc(value: object) -> int:
    text = display_date(value)
    try:
        return -dt.date.fromisoformat(text).toordinal()
    except Exception:
        try:
            return -dt.datetime.strptime(text, "%Y-%m").date().toordinal()
        except Exception:
            return 0


def work_meta_line(article: dict) -> str:
    media = site_media_name(article.get("media"))
    date = display_date_dot(article.get("publish_date"))
    if media and date:
        return f"{media}｜{date}"
    return media or date or ""


def resolve_asset_path(value: object) -> Path | None:
    text = normalize_spaces(value)
    if not text or is_remote_asset(text):
        return None
    raw_path = Path(text).expanduser()
    candidates: list[Path] = []
    if raw_path.is_absolute():
        candidates.append(raw_path)
    else:
        candidates.extend([ROOT / raw_path, ASSETS / raw_path])

    image_exts = [".jpg", ".jpeg", ".png", ".webp", ".avif"]
    for path in list(candidates):
        if path.suffix.lower() in image_exts:
            candidates.extend(path.with_suffix(ext) for ext in image_exts if ext != path.suffix.lower())
        else:
            candidates.extend(path.with_suffix(ext) for ext in image_exts)

    seen = set()
    for path in candidates:
        key = str(path)
        if key in seen:
            continue
        seen.add(key)
        if path.exists() and path.is_file():
            return path
    return None


def site_image_src(value: object, prefix: str = "") -> str:
    text = normalize_spaces(value)
    if not text:
        return ""
    if is_remote_asset(text):
        return text
    source = resolve_asset_path(text)
    if not source:
        return ""
    media_dir = PUBLIC / "assets" / "media"
    media_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha1(str(source).encode("utf-8")).hexdigest()[:8]
    name = f"{slugify(source.stem)}-{digest}{source.suffix.lower()}"
    target = media_dir / name
    if (
        not target.exists()
        or target.stat().st_size != source.stat().st_size
        or target.stat().st_mtime < source.stat().st_mtime
    ):
        shutil.copy2(source, target)
    return f"{prefix}assets/media/{name}"


def picture_frame(class_name: str, src: str, alt: object) -> str:
    safe_src = html_escape(src)
    safe_alt = html_escape(alt)
    return (
        f'<figure class="{class_name}">'
        f'<img class="image-main" src="{safe_src}" alt="{safe_alt}" loading="lazy">'
        "</figure>"
    )


def heading_id(text: str, used: set[str]) -> str:
    base = slugify(display_writing_type(text), "section")
    result = base
    index = 2
    while result in used:
        result = f"{base}-{index}"
        index += 1
    used.add(result)
    return result


def article_body_and_toc(body: str) -> tuple[str, list[tuple[str, str, int]]]:
    blocks = []
    toc: list[tuple[str, str, int]] = []
    used: set[str] = set()
    for block in [item.strip() for item in body.split("\n\n") if item.strip()]:
        heading = re.match(r"^(#{1,6})\s+(.+)$", block)
        if heading and "\n" not in block:
            level = len(heading.group(1))
            tag = "h2" if level <= 2 else "h3" if level == 3 else "h4"
            title = normalize_spaces(heading.group(2))
            anchor = heading_id(title, used)
            toc.append((title, anchor, level))
            blocks.append(f'<{tag} id="{html_escape(anchor)}">{html_escape(title)}</{tag}>')
            continue
        paragraph = "<br>".join(html_escape(line.strip()) for line in block.splitlines() if line.strip())
        if paragraph:
            blocks.append(f"<p>{paragraph}</p>")
    return "".join(blocks), toc


def toc_html(toc: list[tuple[str, str, int]], prefix: str = "") -> str:
    if not toc:
        return '<p class="toc-empty">暂无章节标题</p>'
    links = "\n".join(
        f'<a class="toc-link toc-level-{min(level, 4)}" href="{prefix}#{html_escape(anchor)}">{html_escape(title)}</a>'
        for title, anchor, level in toc
    )
    return f'<nav class="toc-list" aria-label="文章目录">{links}</nav>'


def markdown_body_to_html(body: str) -> str:
    body_html, _ = article_body_and_toc(body)
    return body_html


def first_paragraph(body: str, limit: int = 140) -> str:
    paragraph = next((p.strip() for p in body.split("\n\n") if p.strip()), "")
    return trim_text(paragraph, limit)


def trim_text(text: str, limit: int) -> str:
    clean = normalize_spaces(text)
    if len(clean) <= limit:
        return clean
    return clean[: limit - 1].rstrip() + "…"


TYPE_CANONICAL = {
    "人物专访": "人物与访谈",
    "人物与访谈": "人物与访谈",
    "文旅": "城市、旅行与生活方式",
    "城市、旅行与生活方式": "城市、旅行与生活方式",
    "设计、建筑与文化": "设计、建筑与文化",
    "商业、科技与社会": "商业、科技与社会",
    "品牌稿": "品牌特稿",
    "品牌特稿": "品牌特稿",
}


TYPE_ENGLISH = {
    "人物与访谈": "Profiles & Interviews",
    "城市、旅行与生活方式": "Cities, Travel & Lifestyle",
    "设计、建筑与文化": "Design, Architecture & Culture",
    "商业、科技与社会": "Business, Technology & Society",
    "品牌特稿": "Brand Features",
}


SITE_TYPE_ORDER = [
    "人物与访谈",
    "城市、旅行与生活方式",
    "设计、建筑与文化",
    "商业、科技与社会",
    "品牌特稿",
]


def canonical_writing_type(value: object) -> str:
    display = display_writing_type(value)
    return TYPE_CANONICAL.get(display, display or "未分类")


def type_sort_index(value: object) -> int:
    writing_type = canonical_writing_type(value)
    try:
        return SITE_TYPE_ORDER.index(writing_type)
    except ValueError:
        return len(SITE_TYPE_ORDER) + 1


def ordered_writing_types(articles: list[dict]) -> list[str]:
    present = {canonical_writing_type(article.get("writing_type") or "未分类") for article in articles}
    ordered = [item for item in SITE_TYPE_ORDER if item in present]
    ordered.extend(sorted(present - set(ordered)))
    return ordered


def site_writing_type(value: object) -> str:
    return canonical_writing_type(value)


def ordered_site_writing_types(articles: list[dict]) -> list[str]:
    present = {site_writing_type(article.get("writing_type") or "未分类") for article in articles}
    ordered = [item for item in SITE_TYPE_ORDER if item in present]
    ordered.extend(sorted(present - set(ordered)))
    return ordered


def is_featured_article(article: dict) -> bool:
    return str(article.get("featured") or "").strip() == "是"


def featured_articles(articles: list[dict]) -> list[dict]:
    return sorted(
        [article for article in articles if is_featured_article(article)],
        key=lambda item: (type_sort_index(item.get("writing_type")), date_sort_desc(item.get("publish_date"))),
    )


def media_names(articles: list[dict]) -> list[str]:
    counts: dict[str, int] = {}
    for article in articles:
        name = site_media_name(article.get("media"))
        if not name:
            continue
        counts[name] = counts.get(name, 0) + 1
    return sorted(counts, key=lambda name: (-counts[name], name))


def site_media_name(value: object) -> str:
    name = normalize_spaces(value)
    if name == "Wallpaper":
        return "Wallpaper中文版"
    return name


def read_contact_info() -> dict[str, str]:
    path = INPUT / "contact_info_template.md"
    info = {"email": "", "city": ""}
    if not path.exists():
        return info
    current = ""
    sections: dict[str, list[str]] = {}
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
            continue
        if current and line and not line.startswith("#"):
            sections[current].append(line)
    info["email"] = next(iter(sections.get("Email", [])), "")
    info["city"] = next(iter(sections.get("所在城市", [])), "")
    return info


def latest_full_pdf_path() -> Path | None:
    candidates = sorted(PDF_DIR.glob("*SuFAN全部作品集-摘要版.pdf"), key=lambda path: path.stat().st_mtime, reverse=True)
    if not candidates:
        candidates = sorted(PDF_DIR.glob("*SuFAN全部作品集*.pdf"), key=lambda path: path.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None


def latest_full_pdf_site_link() -> str:
    source = latest_full_pdf_path()
    if not source:
        return ""
    target = PUBLIC / "assets" / source.name
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)
    return f"assets/{source.name}"


def contact_icon_src(filename: str) -> str:
    return site_image_src(filename)


def contact_page_html() -> str:
    info = read_contact_info()
    email = info.get("email") or "待填写"
    city = info.get("city") or "待填写"
    email_href = f"mailto:{email}" if email != "待填写" else "#"
    email_icon = contact_icon_src("email.png")
    wechat_icon = contact_icon_src("微信.png")
    redbook_icon = contact_icon_src("小红书.png")
    quiet_icon = contact_icon_src("MyQuietDraft.jpg")
    pdf_icon = contact_icon_src("pdficon.png")
    pdf_link = latest_full_pdf_site_link()
    pdf_html = (
        f'<a class="social-item" href="{html_escape(pdf_link)}" download><img src="{html_escape(pdf_icon)}" alt="PDF" loading="lazy"><span>下载最新作品集</span></a>'
        if pdf_link
        else '<span class="download-link is-disabled">暂无完整作品集 PDF</span>'
    )
    return f"""
    <section class="page-title contact-title">
      <p>Contact</p>
      <h1>联系方式</h1>
      <div class="intro-copy">欢迎联络洽谈合作。</div>
    </section>
    <section class="contact-grid">
      <div class="contact-block">
        <span>联系方式</span>
        <div class="social-list">
          <a class="social-item" href="{html_escape(email_href)}">
            <img src="{html_escape(email_icon)}" alt="Email" loading="lazy">
            <span>{html_escape(email)}</span>
          </a>
          <button class="social-item copy-social" type="button" data-copy="admin_sue">
            <img src="{html_escape(wechat_icon)}" alt="微信" loading="lazy">
            <span>admin_sue</span>
          </button>
        </div>
      </div>
      <div class="contact-block">
        <span>所在城市</span>
        <p>{html_escape(city)}</p>
      </div>
      <div class="contact-block">
        <span>社交媒体</span>
        <div class="social-list">
          <a class="social-item" href="https://www.xiaohongshu.com/user/profile/5bfe22e351783a3f10144fe4" target="_blank" rel="noopener noreferrer">
            <img src="{html_escape(redbook_icon)}" alt="小红书" loading="lazy">
            <span>小红书 @Susu树树，分享日常写作与作品心得</span>
          </a>
          <a class="social-item" href="https://mp.weixin.qq.com/s/-5Cek0RtiWKjwVZ0lsGXxw" target="_blank" rel="noopener noreferrer">
            <img src="{html_escape(quiet_icon)}" alt="MyQuietDraft" loading="lazy">
            <span>公众号 @MyQuietDraft，分享自由职业与创作体悟</span>
          </a>
        </div>
      </div>
      <div class="contact-block">
        <span>作品集</span>
        <div class="social-list">{pdf_html}</div>
      </div>
    </section>
    """


def build_site() -> None:
    ensure_dirs()
    articles = load_summary_articles(public_only=True)
    if PUBLIC.exists():
        shutil.rmtree(PUBLIC)
    (PUBLIC / "articles").mkdir(parents=True)
    (PUBLIC / "types").mkdir(parents=True)
    (PUBLIC / "assets").mkdir(parents=True)
    write_css()

    types = ordered_site_writing_types(articles)
    media = media_names(articles)
    archive_articles = sorted(articles, key=lambda item: date_sort_desc(item.get("publish_date")))
    hero_image = site_image_src(profile_image_path())
    hero_photo = (
        picture_frame("hero-photo", hero_image, "Su FAN")
        if hero_image
        else ""
    )
    featured_items = featured_articles(articles)

    about_body = about_page_html(articles, types, media, hero_photo)
    write_html(PUBLIC / "index.html", "Su FAN, 自由撰稿人", about_body, active="about")
    write_html(PUBLIC / "about.html", "关于我 - Su FAN", about_body, active="about")

    selected = "\n".join(
        feature_article_card(article)
        for article in featured_items
    )
    write_html(
        PUBLIC / "selected.html",
        "精选作品 - Su FAN",
        f"""
        <section class="page-title selected-title">
          <p>Selected Works</p>
          <h1>精选作品</h1>
          <div class="intro-copy">各领域精选代表作品。</div>
        </section>
        <section class="selected-layout">{selected or '<p class="empty">暂无精选作品。</p>'}</section>
        """,
        active="selected",
    )

    archive_grid = "\n".join(work_grid_card(article) for article in archive_articles) or '<p class="empty">暂无公开作品。</p>'
    archive_list = "\n".join(article_row(article) for article in archive_articles) or '<p class="empty">暂无公开作品。</p>'
    type_filters = "\n".join(
        f'<button class="type-filter" type="button" data-type="{html_escape(t)}">{html_escape(t)}</button>'
        for t in types
    )
    write_html(
        PUBLIC / "archive.html",
        "作品总览 - Su FAN",
        f"""
        <section class="page-title archive-title">
          <p>Archive</p>
          <h1>作品总览</h1>
          <div class="intro-copy">2025年至今完整发布作品总览，含新媒体发布与纸刊发布。</div>
        </section>
        <section class="archive-controls" aria-label="作品筛选与视图切换">
          <div class="type-filter-group">
            <button class="type-filter is-active" type="button" data-type="all">全部</button>
            {type_filters}
          </div>
          <div class="view-toggle" aria-label="浏览模式">
            <button class="view-button is-active" type="button" data-view="grid"><span class="view-icon view-icon-grid" aria-hidden="true"></span><span class="visually-hidden">Grid View</span></button>
            <button class="view-button" type="button" data-view="list"><span class="view-icon view-icon-list" aria-hidden="true"></span><span class="visually-hidden">List View</span></button>
          </div>
        </section>
        <section class="archive-view archive-grid-view is-active" data-view-panel="grid">{archive_grid}</section>
        <section class="archive-view archive-list-view" data-view-panel="list"><div class="archive-list">{archive_list}</div></section>
        """,
        active="archive",
    )

    write_html(PUBLIC / "contact.html", "联系方式 - Su FAN", contact_page_html(), active="contact")

    for writing_type in types:
        items = [article for article in articles if site_writing_type(article.get("writing_type") or "未分类") == writing_type]
        write_html(
            PUBLIC / "types" / f"{slugify(writing_type)}.html",
            f"{writing_type} - 作品集",
            f"""
            <section class="page-title">
              <p>Writing Type</p>
              <h1>{html_escape(writing_type)}</h1>
              <p>{len(items)} 篇公开作品</p>
            </section>
            <main class="archive-list">{''.join(article_row(article, prefix='../') for article in items)}</main>
            """,
            asset_prefix="../",
            active="archive",
        )

    for article in articles:
        body_source = public_body(article)
        keyline = normalize_spaces(article.get("keyline"))
        keyline_html = (
            f'<p class="article-keyline">{html_escape(keyline)}</p>'
            if confirmed_summary(article) and keyline
            else ""
        )
        body_html, toc = article_body_and_toc(body_source)
        tags = " ".join(f"<span>{html_escape(tag)}</span>" for tag in article.get("tags", []))
        article_image = site_image_src(article.get("article_image"), "../")
        article_image_html = (
            picture_frame("article-hero-image", article_image, article.get("title"))
            if article_image
            else ""
        )
        article_class = "article-read has-hero-image" if article_image else "article-read"
        write_html(
            PUBLIC / "articles" / f"{article['id']}.html",
            f"{article.get('title')} - 作品集",
            f"""
            <article class="{article_class}">
              <aside class="article-visual">
                {article_image_html}
                <div class="article-toc">
                  <span>目录</span>
                  {toc_html(toc) if confirmed_summary(article) else '<p class="toc-empty">摘要待确认</p>'}
                </div>
              </aside>
              <div class="article-content">
                <a class="back-link js-back-link" href="../archive.html">返回上页</a>
                <header class="article-header">
                  <p>Published Work</p>
                  <h1>{html_escape(article.get('title'))}</h1>
                  <div class="meta">
                    <span>{html_escape(site_media_name(article.get('media')))}</span>
                    <span>{html_escape(article.get('publish_date'))}</span>
                    <span>{html_escape(site_writing_type(article.get('writing_type')))}</span>
                    {tags}
                  </div>
                </header>
                {keyline_html}
                <div class="article-divider"></div>
                {article_source_prompt_html(article)}
                <div class="body">{body_html if confirmed_summary(article) else '<p>摘要待确认。请点击阅读原文查看完整作品。</p>'}</div>
              </div>
            </article>
            """,
            asset_prefix="../",
            active="archive",
        )

    write_static_site_metadata()
    print(f"网站已更新：{PUBLIC / 'index.html'}")


def write_static_site_metadata() -> None:
    PUBLIC.mkdir(parents=True, exist_ok=True)
    (PUBLIC / ".nojekyll").write_text("", encoding="utf-8")


def load_env_file(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    if not path.exists():
        return values
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def config_value(config: dict[str, str], key: str, default: str = "") -> str:
    return os.environ.get(key, config.get(key, default)).strip()


def path_inside(child: Path, parent: Path) -> bool:
    try:
        child.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def site_footer_html() -> str:
    text = config_value(load_env_file(TENCENT_ENV), "ICP_BEIAN_TEXT")
    beian = (
        f'<a href="https://beian.miit.gov.cn/" target="_blank" rel="noopener">{html_escape(text)}</a>'
        if text
        else ""
    )
    return (
        '<footer class="site-footer">'
        '<span class="site-copyright">© 2026 Su FAN. All Rights Reserved.</span>'
        f'<span class="site-beian">{beian}</span>'
        "</footer>"
    )


def run_git(*args: str, check: bool = True) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        ["git", *args],
        cwd=ROOT,
        check=check,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def is_git_repo() -> bool:
    result = run_git("rev-parse", "--is-inside-work-tree", check=False)
    return result.returncode == 0 and result.stdout.strip() == "true"


def ensure_git_repo() -> None:
    if not is_git_repo():
        subprocess.run(["git", "init"], cwd=ROOT, check=True)
    head = run_git("symbolic-ref", "--short", "HEAD", check=False).stdout.strip()
    has_commit = run_git("rev-parse", "--verify", "HEAD", check=False).returncode == 0
    if not has_commit and head != "main":
        subprocess.run(["git", "symbolic-ref", "HEAD", "refs/heads/main"], cwd=ROOT, check=True)


def git_has_changes() -> bool:
    result = run_git("status", "--porcelain")
    return bool(result.stdout.strip())


def current_git_branch() -> str:
    result = run_git("branch", "--show-current", check=False)
    branch = result.stdout.strip()
    return branch or "main"


def has_origin_remote() -> bool:
    result = run_git("remote", "get-url", "origin", check=False)
    return result.returncode == 0 and bool(result.stdout.strip())


def publish_site(message: str) -> None:
    if not message.strip():
        raise SystemExit("请提供本次备份说明，例如：./portfolio.sh publish \"更新作品集网站\"")

    build_site()
    index = PUBLIC / "index.html"
    if not index.exists():
        raise SystemExit("网站生成失败：未找到 output/site/index.html。")

    ensure_git_repo()
    run_git("add", ".")

    if git_has_changes():
        run_git("commit", "-m", message.strip())
        print(f"已提交本次更新：{message.strip()}")
    else:
        print("没有检测到需要提交的变更。")

    if has_origin_remote():
        branch = current_git_branch()
        subprocess.run(["git", "push", "-u", "origin", branch], cwd=ROOT, check=True)
        print("已推送到 GitHub 备份仓库。线上网站请继续运行 Vercel 或腾讯云部署步骤。")
    else:
        print("尚未配置 GitHub 远程仓库。创建公开仓库后运行：")
        print("git remote add origin <你的 GitHub 仓库地址>")
        print(f"git push -u origin {current_git_branch()}")


def deploy_vercel() -> None:
    build_site()
    index = PUBLIC / "index.html"
    if not index.exists():
        raise SystemExit("网站生成失败：未找到 output/site/index.html。")

    vercel_output = ROOT / ".vercel" / "output"
    static_output = vercel_output / "static"
    if static_output.exists():
        shutil.rmtree(static_output)
    static_output.mkdir(parents=True, exist_ok=True)
    shutil.copytree(PUBLIC, static_output, dirs_exist_ok=True, ignore=shutil.ignore_patterns(".vercel", "output"))
    (vercel_output / "config.json").write_text('{"version":3}\n', encoding="utf-8")

    env = os.environ.copy()
    if NODE.exists():
        env["PATH"] = f"{NODE.parent}:{env.get('PATH', '')}"

    npx = shutil.which("npx")
    if npx:
        command = [npx, "vercel", "deploy", "--prebuilt", "--prod"]
    elif PNPM.exists():
        command = [str(PNPM), "dlx", "vercel", "deploy", "--prebuilt", "--prod"]
    else:
        raise SystemExit("找不到 Vercel CLI。也可以进入 Vercel Drop 手动上传 output/site/ 文件夹。")

    subprocess.run(command, cwd=ROOT, env=env, check=True)


def deploy_tencent_server() -> None:
    config = load_env_file(TENCENT_ENV)
    host = config_value(config, "TENCENT_SERVER_HOST")
    user = config_value(config, "TENCENT_SERVER_USER", "root")
    port = config_value(config, "TENCENT_SERVER_PORT", "22")
    remote_path = config_value(config, "TENCENT_SERVER_PATH", "/var/www/sufan-portfolio")
    ssh_key = config_value(config, "TENCENT_SERVER_SSH_KEY")

    missing = [
        name
        for name, value in {
            "TENCENT_SERVER_HOST": host,
            "TENCENT_SERVER_USER": user,
            "TENCENT_SERVER_PORT": port,
            "TENCENT_SERVER_PATH": remote_path,
        }.items()
        if not value
    ]
    if missing:
        raise SystemExit(
            "缺少腾讯云服务器配置："
            + "、".join(missing)
            + "。请复制 .env.tencent-server.example 为 .env.tencent-server 后填写。"
        )

    rsync = shutil.which("rsync")
    ssh = shutil.which("ssh")
    if not rsync or not ssh:
        raise SystemExit("找不到 rsync 或 ssh，无法自动部署到腾讯云服务器。")

    build_site()
    index = PUBLIC / "index.html"
    if not index.exists():
        raise SystemExit("网站生成失败：未找到 output/site/index.html。")

    if ssh_key and path_inside(Path(ssh_key).expanduser(), ROOT):
        raise SystemExit("请不要把 SSH 私钥放在项目文件夹内；建议放在 ~/.ssh/ 并只在 .env.tencent-server 中填写路径。")

    ssh_args = [
        ssh,
        "-p",
        port,
        "-o",
        "BatchMode=yes",
        "-o",
        "PasswordAuthentication=no",
        "-o",
        "PubkeyAuthentication=yes",
        "-o",
        "IdentitiesOnly=yes",
        "-o",
        "StrictHostKeyChecking=accept-new",
    ]
    if ssh_key:
        ssh_args.extend(["-i", ssh_key])
    remote = f"{user}@{host}"
    quoted_path = shlex.quote(remote_path)
    quoted_user = shlex.quote(user)

    subprocess.run(
        [
            *ssh_args,
            remote,
            f"sudo mkdir -p {quoted_path} && sudo chown -R {quoted_user}:{quoted_user} {quoted_path}",
        ],
        cwd=ROOT,
        check=True,
    )

    ssh_command = " ".join(shlex.quote(part) for part in ssh_args)
    subprocess.run(
        [
            rsync,
            "-az",
            "--delete",
            "-e",
            ssh_command,
            f"{PUBLIC}/",
            f"{remote}:{remote_path.rstrip('/')}/",
        ],
        cwd=ROOT,
        check=True,
    )

    subprocess.run(
        [
            *ssh_args,
            remote,
            "sudo nginx -t && sudo systemctl reload nginx",
        ],
        cwd=ROOT,
        check=True,
    )
    print(f"腾讯云服务器部署完成：http://{host}/")


def site_nav(asset_prefix: str, active: str) -> str:
    items = [
        ("about", "关于我", f"{asset_prefix}about.html"),
        ("selected", "精选作品", f"{asset_prefix}selected.html"),
        ("archive", "作品总览", f"{asset_prefix}archive.html"),
        ("contact", "联系方式", f"{asset_prefix}contact.html"),
    ]
    links = "\n".join(
        f'<a class="nav-link{" is-active" if key == active else ""}" href="{href}">{label}</a>'
        for key, label, href in items
    )
    return f"""
    <header class="site-header">
      <a class="brand" href="{asset_prefix}index.html">Su FAN, 自由撰稿人</a>
      <nav class="main-nav" aria-label="主导航">
        {links}
        <button class="theme-toggle" type="button" aria-label="切换明暗模式" aria-pressed="false">
          <span class="theme-option theme-option-sun" aria-hidden="true">
            <span class="theme-icon">☀︎</span>
          </span>
          <span class="theme-option theme-option-moon" aria-hidden="true">
            <span class="theme-icon">☽</span>
          </span>
          <span class="visually-hidden">切换明暗模式</span>
        </button>
      </nav>
    </header>
    """


def site_script() -> str:
    return """
    <script>
    (() => {
      const root = document.documentElement;
      const stored = localStorage.getItem("portfolio-theme");
      const prefersDark = window.matchMedia && window.matchMedia("(prefers-color-scheme: dark)").matches;
      const initial = stored || (prefersDark ? "dark" : "light");
      const button = document.querySelector(".theme-toggle");
      let transitionTimer = 0;
      const setTheme = (theme, animate = false) => {
        if (animate) {
          window.clearTimeout(transitionTimer);
          root.classList.add("theme-transition");
          transitionTimer = window.setTimeout(() => root.classList.remove("theme-transition"), 520);
        }
        root.dataset.theme = theme;
        if (button) {
          button.dataset.theme = theme;
          button.setAttribute("aria-label", theme === "dark" ? "切换到 Light Mode" : "切换到 Dark Mode");
          button.setAttribute("aria-pressed", theme === "dark" ? "true" : "false");
        }
      };
      setTheme(initial);
      button?.addEventListener("click", () => {
        const next = root.dataset.theme === "dark" ? "light" : "dark";
        localStorage.setItem("portfolio-theme", next);
        setTheme(next, true);
      });

      document.querySelectorAll(".copy-social").forEach((copyButton) => {
        copyButton.addEventListener("click", async () => {
          const text = copyButton.dataset.copy || "";
          if (!text) return;
          try {
            await navigator.clipboard.writeText(text);
            copyButton.dataset.copied = "true";
            window.setTimeout(() => {
              copyButton.dataset.copied = "false";
            }, 1400);
          } catch (error) {
            copyButton.dataset.copied = "false";
          }
        });
      });

      document.querySelectorAll(".js-back-link").forEach((link) => {
        link.addEventListener("click", (event) => {
          if (document.referrer && new URL(document.referrer).origin === window.location.origin && window.history.length > 1) {
            event.preventDefault();
            window.history.back();
          }
        });
      });

      const archive = document.querySelector(".archive-title");
      if (!archive) return;
      const typeButtons = Array.from(document.querySelectorAll(".type-filter"));
      const viewButtons = Array.from(document.querySelectorAll(".view-button"));
      const panels = Array.from(document.querySelectorAll("[data-view-panel]"));
      const items = Array.from(document.querySelectorAll("[data-work-type]"));

      typeButtons.forEach((button) => {
        button.addEventListener("click", () => {
          const type = button.dataset.type || "all";
          typeButtons.forEach((item) => item.classList.toggle("is-active", item === button));
          items.forEach((item) => {
            item.hidden = type !== "all" && item.dataset.workType !== type;
          });
          const url = new URL(window.location.href);
          if (type === "all") {
            url.searchParams.delete("type");
          } else {
            url.searchParams.set("type", type);
          }
          window.history.replaceState({}, "", url);
        });
      });

      viewButtons.forEach((button) => {
        button.addEventListener("click", () => {
          const view = button.dataset.view || "grid";
          viewButtons.forEach((item) => item.classList.toggle("is-active", item === button));
          panels.forEach((panel) => panel.classList.toggle("is-active", panel.dataset.viewPanel === view));
        });
      });

      const initialType = new URLSearchParams(window.location.search).get("type");
      if (initialType) {
        const target = typeButtons.find((button) => button.dataset.type === initialType);
        target?.click();
      }
    })();
    </script>
    """


def about_page_html(articles: list[dict], types: list[str], media: list[str], hero_photo: str) -> str:
    type_blocks = "\n".join(
        f"""
        <a class="field-card" href="archive.html?type={urllib.parse.quote(writing_type)}" data-field="{html_escape(writing_type)}">
          <span class="field-cn">{html_escape(writing_type)}</span>
          <span class="field-en">{html_escape(TYPE_ENGLISH.get(writing_type, ""))}</span>
          <p>{sum(1 for article in articles if site_writing_type(article.get('writing_type')) == writing_type)} 篇作品</p>
        </a>
        """
        for writing_type in types
    )
    media_list = " ".join(f"<span>{html_escape(name)}</span>" for name in media)
    return f"""
    <section class="about-hero">
      <div class="about-copy">
        <p>Biography</p>
        <h1>Su FAN 作品集</h1>
        <ul class="hero-bio">
          <li>现为自由撰稿人，base 北京。</li>
          <li>中英文流利，能从容应对双语采访与撰稿任务。</li>
          <li>教育背景：经济学博士，北京大学（2025）；联合培养博士研究生，巴黎第一大学（2023）；经济学、社会学学士，北京大学（2020）。</li>
        </ul>
      </div>
      <div class="about-portrait">{hero_photo}</div>
    </section>
    <section class="about-dashboard" aria-label="作品统计">
      <div><strong>{len(articles)}</strong><span>公开发表作品</span></div>
      <div><strong>{len(media)}</strong><span>合作媒体</span></div>
      <div><strong>{len(types)}</strong><span>写作领域</span></div>
    </section>
    <section class="about-section media-section">
      <h2>合作媒体</h2>
      <div class="media-list">{media_list or '<span>暂无</span>'}</div>
    </section>
    <section class="about-section fields-section">
      <h2>写作领域</h2>
      <div class="field-grid">{type_blocks}</div>
    </section>
    """


def article_image_markup(article: dict, prefix: str, class_name: str) -> str:
    image = site_image_src(article.get("article_image"), prefix)
    if image:
        return picture_frame(class_name, image, article.get("title"))
    return f'<div class="{class_name} image-placeholder" aria-hidden="true"><span>Image pending</span></div>'


def feature_article_card(article: dict, prefix: str = "", featured: bool = False) -> str:
    classes = ["feature-card"]
    if featured:
        classes.append("feature-primary")
    image_html = article_image_markup(article, prefix, "feature-image")
    preview = public_preview(article, 220)
    return f"""
    <article class="{' '.join(classes)}" data-work-type="{html_escape(site_writing_type(article.get('writing_type')))}">
      <a href="{prefix}articles/{html_escape(article['id'])}.html">{image_html}</a>
      <div class="feature-copy">
        <p class="card-type">{html_escape(site_writing_type(article.get('writing_type')))}</p>
        <h2><a href="{prefix}articles/{html_escape(article['id'])}.html">{html_escape(article.get('title'))}</a></h2>
        <p class="card-meta-line">{html_escape(work_meta_line(article))}</p>
        <p class="summary">{html_escape(preview)}</p>
      </div>
    </article>
    """


def work_grid_card(article: dict, prefix: str = "") -> str:
    return f"""
    <article class="work-card" data-work-type="{html_escape(site_writing_type(article.get('writing_type')))}">
      <a class="work-card-image" href="{prefix}articles/{html_escape(article['id'])}.html">
        {article_image_markup(article, prefix, "grid-image")}
      </a>
      <div class="work-card-copy">
        <p class="card-type">{html_escape(site_writing_type(article.get('writing_type')))}</p>
        <h2><a href="{prefix}articles/{html_escape(article['id'])}.html">{html_escape(article.get('title'))}</a></h2>
        <p class="card-meta-line">{html_escape(work_meta_line(article))}</p>
      </div>
    </article>
    """


def article_card(article: dict, prefix: str = "", featured: bool = False) -> str:
    tags = " ".join(f"<span>{html_escape(tag)}</span>" for tag in article.get("tags", []))
    image = site_image_src(article.get("article_image"), prefix)
    classes = ["feature-card"]
    if image:
        classes.append("has-image")
    if featured:
        classes.append("feature-primary")
    class_name = " ".join(classes)
    image_html = (
        picture_frame("feature-image", image, article.get("title"))
        if image
        else ""
    )
    preview = public_preview(article, 240)
    return f"""
    <article class="{class_name}">
      <div class="feature-copy">
        <div class="card-kicker">
          <span>{html_escape(site_writing_type(article.get('writing_type')))}</span>
          <span>{html_escape(article.get('publish_date'))}</span>
        </div>
        <h2><a href="{prefix}articles/{html_escape(article['id'])}.html">{html_escape(article.get('title'))}</a></h2>
        {source_link_html(article)}
        <p class="summary">{html_escape(preview)}</p>
        <div class="meta">
          <span>{html_escape(site_media_name(article.get('media')))}</span>
          {tags}
        </div>
      </div>
      {image_html}
    </article>
    """


def article_row(article: dict, prefix: str = "") -> str:
    image = site_image_src(article.get("article_image"), prefix)
    image_html = (
        picture_frame("archive-image", image, article.get("title"))
        if image
        else '<div class="archive-image archive-image-placeholder" aria-hidden="true"></div>'
    )
    return f"""
    <article class="archive-item" data-work-type="{html_escape(site_writing_type(article.get('writing_type')))}">
      {image_html}
      <div class="archive-main">
        <p class="card-type">{html_escape(site_writing_type(article.get('writing_type')))}</p>
        <h3><a href="{prefix}articles/{html_escape(article['id'])}.html">{html_escape(article.get('title'))}</a></h3>
        <p class="card-meta-line">{html_escape(work_meta_line(article))}</p>
        {source_link_html(article)}
      </div>
    </article>
    """


def write_html(path: Path, title: str, body: str, asset_prefix: str = "", active: str = "") -> None:
    footer = site_footer_html()
    page_class = f"page-{path.stem}"
    path.write_text(
        textwrap.dedent(
            f"""\
            <!doctype html>
            <html lang="zh-CN">
            <head>
              <meta charset="utf-8">
              <meta name="viewport" content="width=device-width, initial-scale=1">
              <title>{html_escape(title)}</title>
              <link rel="stylesheet" href="{asset_prefix}assets/site.css">
              <script>
                (() => {{
                  const stored = localStorage.getItem("portfolio-theme");
                  const prefersDark = window.matchMedia && window.matchMedia("(prefers-color-scheme: dark)").matches;
                  document.documentElement.dataset.theme = stored || (prefersDark ? "dark" : "light");
                }})();
              </script>
            </head>
            <body class="{html_escape(page_class)}">
              {textwrap.dedent(site_nav(asset_prefix, active)).strip()}
              <div class="shell">
                {textwrap.dedent(body).strip()}
              </div>
              {footer}
              {textwrap.dedent(site_script()).strip()}
            </body>
            </html>
            """
        ),
        encoding="utf-8",
    )


def write_css() -> None:
    (PUBLIC / "assets" / "site.css").write_text(
        """
        :root {
          color-scheme: light;
          --paper: #fdfdfb;
          --panel: #f7f7f3;
          --panel-strong: #eeeeea;
          --ink: #121212;
          --muted: #62625d;
          --faint: #969690;
          --line: #deded7;
          --line-strong: #161616;
          --image-soft: rgba(18, 18, 18, 0.08);
          --accent: #121212;
          --switch-bg: #f2f2ef;
          --switch-active: #ffffff;
          --switch-active-ink: #121212;
          --switch-shadow: rgba(18, 18, 18, 0.1);
          --radius: 18px;
          --ease: cubic-bezier(0.16, 1, 0.3, 1);
          font-family: "Helvetica Neue", Arial, "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
        }
        :root[data-theme="dark"] {
          color-scheme: dark;
          --paper: #11110f;
          --panel: #191916;
          --panel-strong: #24241f;
          --ink: #f2f2ee;
          --muted: #b6b6ad;
          --faint: #7f7f77;
          --line: #34342e;
          --line-strong: #f2f2ee;
          --image-soft: rgba(242, 242, 238, 0.12);
          --accent: #f2f2ee;
          --switch-bg: #24241f;
          --switch-active: #f2f2ee;
          --switch-active-ink: #11110f;
          --switch-shadow: rgba(0, 0, 0, 0.28);
        }
        * { box-sizing: border-box; }
        html {
          scroll-behavior: smooth;
          background: var(--paper);
        }
        body {
          margin: 0;
          background: var(--paper);
          color: var(--ink);
          font-size: 16px;
          line-height: 1.7;
          text-rendering: optimizeLegibility;
          -webkit-font-smoothing: antialiased;
        }
        a { color: inherit; text-decoration: none; }
        button { font: inherit; }
        [hidden] { display: none !important; }
        a:focus-visible {
          outline: 1px solid var(--ink);
          outline-offset: 4px;
        }
        button:focus-visible {
          outline: 1px solid var(--ink);
          outline-offset: 4px;
        }
        @media (prefers-reduced-motion: no-preference) {
          .theme-transition,
          .theme-transition body,
          .theme-transition .site-header,
          .theme-transition .site-footer,
          .theme-transition .shell,
          .theme-transition .theme-toggle,
          .theme-transition .theme-option,
          .theme-transition .feature-card,
          .theme-transition .work-card,
          .theme-transition .archive-item,
          .theme-transition .about-dashboard,
          .theme-transition .field-card,
          .theme-transition .contact-block,
          .theme-transition .article-layout,
          .theme-transition .article-visual,
          .theme-transition .article-content,
          .theme-transition .body,
          .theme-transition .toc-panel,
          .theme-transition a,
          .theme-transition button {
            transition:
              background-color 420ms var(--ease),
              color 420ms var(--ease),
              border-color 420ms var(--ease),
              box-shadow 420ms var(--ease),
              opacity 420ms var(--ease);
          }
        }
        @keyframes editorial-rise {
          from {
            opacity: 0;
            transform: translateY(12px);
          }
          to {
            opacity: 1;
            transform: translateY(0);
          }
        }
        .shell > *,
        .feature-card,
        .work-card,
        .archive-item,
        .field-card,
        .contact-block,
        .article-visual,
        .article-content {
          animation: editorial-rise 640ms var(--ease) both;
        }
        .shell > *:nth-child(2) { animation-delay: 70ms; }
        .shell > *:nth-child(3) { animation-delay: 120ms; }
        .feature-card:nth-child(2),
        .work-card:nth-child(2),
        .archive-item:nth-child(2),
        .field-card:nth-child(2),
        .contact-block:nth-child(2) { animation-delay: 80ms; }
        .feature-card:nth-child(3),
        .work-card:nth-child(3),
        .archive-item:nth-child(3),
        .field-card:nth-child(3),
        .contact-block:nth-child(3) { animation-delay: 130ms; }
        .feature-card:nth-child(n+4),
        .work-card:nth-child(n+4),
        .archive-item:nth-child(n+4),
        .field-card:nth-child(n+4),
        .contact-block:nth-child(n+4) { animation-delay: 170ms; }
        .site-header {
          position: sticky;
          top: 0;
          z-index: 10;
          display: flex;
          align-items: center;
          justify-content: space-between;
          gap: 24px;
          width: min(1180px, calc(100% - 48px));
          min-height: 68px;
          margin: 0 auto;
          padding: 0;
          border-bottom: 1px solid var(--line);
          background: color-mix(in srgb, var(--paper) 92%, transparent);
          backdrop-filter: blur(18px);
          -webkit-backdrop-filter: blur(18px);
        }
        .brand {
          flex: 0 0 auto;
          color: var(--ink);
          font-size: 15px;
          font-weight: 500;
          line-height: 1;
          white-space: nowrap;
        }
        .main-nav {
          display: flex;
          align-items: center;
          justify-content: flex-end;
          gap: 20px;
          min-width: 0;
        }
        .nav-link,
        .theme-toggle,
        .view-button,
        .type-filter {
          border: 0;
          border-bottom: 1px solid transparent;
          padding: 4px 0;
          background: transparent;
          color: var(--muted);
          cursor: pointer;
          font-size: 14px;
          line-height: 1.3;
          white-space: nowrap;
          transition: color 180ms var(--ease), border-color 180ms var(--ease);
        }
        .nav-link:hover,
        .nav-link.is-active,
        .theme-toggle:hover,
        .view-button:hover,
        .view-button.is-active,
        .type-filter:hover,
        .type-filter.is-active {
          border-color: currentColor;
          color: var(--ink);
        }
        .theme-toggle {
          display: inline-flex;
          align-items: center;
          gap: 2px;
          width: 74px;
          height: 36px;
          border: 0;
          border-radius: 999px;
          padding: 3px;
          background: var(--switch-bg);
          color: var(--faint);
          box-shadow: inset 0 0 0 1px color-mix(in srgb, var(--ink) 5%, transparent);
          transition: transform 180ms var(--ease), box-shadow 180ms var(--ease);
        }
        .theme-toggle:hover {
          transform: translateY(-1px);
          box-shadow: inset 0 0 0 1px color-mix(in srgb, var(--ink) 10%, transparent);
        }
        .theme-option {
          position: relative;
          display: grid;
          width: 30px;
          height: 30px;
          place-items: center;
          border-radius: 999px;
          background: transparent;
          color: currentColor;
          transition: background-color 180ms var(--ease), color 180ms var(--ease), box-shadow 180ms var(--ease);
        }
        .theme-toggle[data-theme="light"] .theme-option-sun,
        .theme-toggle[data-theme="dark"] .theme-option-moon {
          background: var(--switch-active);
          color: var(--switch-active-ink);
          box-shadow: 0 2px 10px var(--switch-shadow);
        }
        .theme-icon {
          display: block;
          font-size: 19px;
          font-weight: 400;
          line-height: 1;
        }
        .visually-hidden {
          position: absolute;
          width: 1px;
          height: 1px;
          padding: 0;
          margin: -1px;
          overflow: hidden;
          clip: rect(0, 0, 0, 0);
          white-space: nowrap;
          border: 0;
        }
        .shell {
          width: min(1180px, calc(100% - 48px));
          margin: 0 auto;
          padding: 56px 0 108px;
        }
        .site-footer {
          width: min(1180px, calc(100% - 48px));
          margin: -72px auto 40px;
          padding-top: 18px;
          border-top: 1px solid var(--line);
          color: var(--muted);
          font-size: 13px;
          line-height: 1.6;
          display: flex;
          align-items: center;
          justify-content: space-between;
          gap: 24px;
        }
        .site-footer a {
          color: inherit;
          text-decoration: none;
        }
        .site-footer a:hover {
          color: var(--ink);
        }
        .page-title,
        .about-hero {
          position: relative;
          padding: 64px 0 44px;
          border-bottom: 1px solid var(--line);
        }
        .page-title h1,
        .about-hero h1,
        .article-header h1 {
          max-width: 920px;
          margin: 0.24em 0 0.48em;
          color: var(--ink);
          font-size: 42px;
          font-weight: 300;
          line-height: 1.12;
          letter-spacing: 0;
        }
        .page-title > p,
        .about-copy > p,
        .article-header > p {
          margin: 0 0 12px;
          color: var(--muted);
          font-size: 13px;
          font-weight: 500;
          line-height: 1.35;
          text-transform: none;
        }
        .intro-copy {
          max-width: 680px;
          margin: 0;
          color: var(--muted);
          font-size: 16px;
          font-weight: 400;
          line-height: 1.75;
        }
        .about-hero {
          display: grid;
          grid-template-columns: minmax(0, 1.15fr) minmax(360px, 0.85fr);
          gap: 54px;
          align-items: start;
        }
        .about-portrait {
          display: grid;
          justify-items: end;
        }
        .hero-photo {
          width: min(330px, 100%);
        }
        .hero-photo,
        .feature-image,
        .archive-image,
        .grid-image,
        .article-hero-image {
          position: relative;
          margin: 0;
          overflow: hidden;
          background: var(--panel);
          border-radius: var(--radius);
          isolation: isolate;
        }
        .image-main {
          display: block;
          width: 100%;
          height: auto;
          transform-origin: center;
          transition: transform 420ms var(--ease);
        }
        a:hover .image-main,
        .feature-image:hover .image-main,
        .archive-image:hover .image-main,
        .grid-image:hover .image-main,
        .article-hero-image:hover .image-main {
          transform: scale(1.035);
        }
        .hero-bio {
          margin: 0;
          padding: 0;
          list-style: none;
          color: var(--ink);
          font-size: 16px;
          font-weight: 300;
          line-height: 1.9;
          letter-spacing: 0.01em;
        }
        .hero-bio li {
          max-width: 760px;
          margin: 0 0 0.55em;
        }
        .about-dashboard {
          display: grid;
          grid-template-columns: repeat(3, minmax(0, 1fr));
          gap: 24px;
          padding: 24px 0;
          border-top: 1px solid var(--line);
          border-bottom: 1px solid var(--line);
        }
        .about-dashboard div {
          padding: 0 0 18px;
          border-bottom: 1px solid var(--line);
        }
        .page-about .about-dashboard div {
          padding-bottom: 0;
          border-bottom: 0;
        }
        .about-dashboard strong {
          display: block;
          color: var(--ink);
          font-size: 36px;
          font-weight: 300;
          line-height: 1;
        }
        .about-dashboard span {
          display: block;
          margin-top: 8px;
          color: var(--muted);
          font-size: 13px;
          line-height: 1.4;
        }
        .about-section {
          padding: 58px 0;
          border-bottom: 1px solid var(--line);
        }
        .page-about .about-section:last-child {
          border-bottom: 0;
        }
        .about-section h2 {
          margin: 0 0 24px;
          color: var(--ink);
          font-size: 28px;
          font-weight: 300;
          line-height: 1.2;
        }
        .media-list {
          display: flex;
          flex-wrap: wrap;
          gap: 10px;
        }
        .media-list span {
          border: 1px solid var(--line);
          border-radius: 999px;
          padding: 7px 12px;
          color: var(--muted);
          font-size: 13px;
          font-weight: 300;
          line-height: 1.4;
        }
        .field-grid {
          display: grid;
          grid-template-columns: repeat(3, minmax(0, 1fr));
          gap: 14px;
        }
        .field-card {
          min-height: 150px;
          border: 1px solid var(--line);
          border-radius: var(--radius);
          padding: 20px;
          background: var(--panel);
          transition: border-color 180ms var(--ease), transform 180ms var(--ease), background 180ms var(--ease);
        }
        .field-card:hover {
          border-color: var(--ink);
          transform: translateY(-2px);
        }
        .field-card .field-cn {
          display: block;
          color: var(--ink);
          font-size: 19px;
          font-weight: 300;
          line-height: 1.35;
        }
        .field-card .field-en {
          display: block;
          margin-top: 8px;
          color: var(--muted);
          font-size: 12px;
          font-weight: 400;
          line-height: 1.35;
        }
        .field-card p {
          margin: 30px 0 0;
          color: var(--muted);
          font-size: 13px;
        }
        .source-link {
          margin: 0.3em 0 1em;
          color: var(--muted);
          font-size: 13px;
          font-weight: 400;
          line-height: 1.5;
        }
        .source-link a {
          border-bottom: 1px solid currentColor;
        }
        .source-link a:hover {
          color: var(--ink);
        }
        .selected-layout {
          display: grid;
          grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
          gap: 38px 26px;
          padding-top: 42px;
        }
        .feature-card {
          display: grid;
          grid-template-rows: auto 1fr;
          gap: 18px;
          align-content: start;
        }
        .feature-copy,
        .work-card-copy,
        .archive-main {
          display: grid;
          grid-template-rows: auto minmax(3.4em, auto) auto auto;
          align-content: start;
        }
        .card-type {
          margin: 0 0 8px;
          color: var(--muted);
          font-size: 13px;
          font-weight: 400;
          line-height: 1.4;
        }
        .feature-card h2 {
          max-width: 760px;
          margin: 0;
          color: var(--ink);
          font-size: 21px;
          font-weight: 300;
          line-height: 1.3;
          letter-spacing: 0;
        }
        .card-meta-line {
          margin: 10px 0 0;
          color: var(--muted);
          font-size: 13px;
          font-weight: 400;
          line-height: 1.45;
        }
        .summary {
          color: var(--ink);
          max-width: 760px;
          font-size: 16px;
          font-weight: 400;
          line-height: 1.9;
          margin: 1em 0;
        }
        .selected-layout .summary {
          color: var(--muted);
          font-size: 14px;
          line-height: 1.75;
        }
        .archive-controls {
          display: grid;
          grid-template-columns: minmax(0, 1fr) auto;
          gap: 28px;
          align-items: start;
          padding: 28px 0;
          border-bottom: 1px solid var(--line);
        }
        .type-filter-group {
          display: flex;
          flex-wrap: wrap;
          gap: 16px 22px;
        }
        .view-toggle {
          display: flex;
          gap: 12px;
          align-items: center;
        }
        .view-button {
          display: inline-grid;
          width: 32px;
          height: 32px;
          place-items: center;
          border: 1px solid transparent;
          border-radius: 999px;
          padding: 0;
        }
        .view-button:hover,
        .view-button.is-active {
          border-color: var(--line);
          background: var(--panel);
        }
        .view-icon {
          display: block;
          width: 17px;
          height: 17px;
          color: currentColor;
        }
        .view-icon-grid {
          background:
            radial-gradient(currentColor 1.6px, transparent 1.8px) 0 0 / 33.333% 33.333%;
        }
        .view-icon-list {
          position: relative;
        }
        .view-icon-list::before {
          content: "";
          position: absolute;
          left: 0;
          right: 0;
          top: 3px;
          height: 1px;
          background: currentColor;
          box-shadow: 0 5px 0 currentColor, 0 10px 0 currentColor;
        }
        .archive-view { display: none; }
        .archive-view.is-active { display: block; }
        .archive-grid-view.is-active {
          display: grid;
          grid-template-columns: repeat(auto-fill, minmax(190px, 1fr));
          gap: 34px 22px;
          padding-top: 36px;
        }
        .work-card {
          display: grid;
          grid-template-rows: auto 1fr;
          gap: 14px;
          align-content: start;
        }
        .work-card-copy h2 {
          margin: 0;
          color: var(--ink);
          font-size: 17px;
          font-weight: 300;
          line-height: 1.35;
        }
        .work-meta,
        .meta {
          display: flex;
          flex-wrap: wrap;
          gap: 8px 12px;
          margin-top: 12px;
          color: var(--muted);
          font-size: 13px;
          font-weight: 400;
          line-height: 1.45;
        }
        .archive-list {
          display: grid;
          border-top: 1px solid var(--line);
        }
        .archive-item {
          display: grid;
          grid-template-columns: 132px minmax(0, 1fr);
          gap: 30px;
          padding: 24px 0;
          border-bottom: 1px solid var(--line);
          align-items: start;
          transition: border-color 160ms var(--ease), background 160ms var(--ease);
        }
        .archive-image {
          width: 100%;
          aspect-ratio: 3 / 4;
          display: grid;
          place-items: center;
        }
        .grid-image,
        .feature-image {
          aspect-ratio: 3 / 4;
          display: grid;
          place-items: center;
        }
        .grid-image .image-main,
        .feature-image .image-main,
        .archive-image .image-main {
          width: 100%;
          height: 100%;
          object-fit: contain;
        }
        .grid-image,
        .feature-image,
        .archive-image,
        .article-hero-image { box-shadow: 0 18px 54px var(--image-soft); }
        .archive-image-placeholder {
          display: grid;
          place-items: center;
          color: var(--faint);
          background: var(--panel);
        }
        .image-placeholder {
          display: grid;
          place-items: center;
          border-radius: var(--radius);
          background: var(--panel-strong);
          color: var(--faint);
          font-size: 13px;
        }
        .archive-main h3 {
          max-width: 880px;
          margin: 0;
          color: var(--ink);
          font-size: 18px;
          font-weight: 300;
          line-height: 1.35;
          letter-spacing: 0;
        }
        .back-link {
          display: inline-block;
          margin-bottom: 32px;
          color: var(--muted);
          font-size: 13px;
          border-bottom: 1px solid currentColor;
        }
        .article-read {
          display: grid;
          grid-template-columns: minmax(280px, 400px) minmax(0, 720px);
          gap: 64px;
          align-items: start;
        }
        .article-visual {
          position: sticky;
          top: 92px;
          display: grid;
          gap: 20px;
          max-height: calc(100dvh - 120px);
          overflow: auto;
          padding-bottom: 4px;
        }
        .article-hero-image {
          position: relative;
          overflow: hidden;
          border-radius: var(--radius);
          background: var(--panel);
        }
        .article-toc {
          padding-top: 18px;
          border-top: 1px solid var(--line);
        }
        .article-toc > span {
          display: block;
          margin-bottom: 12px;
          color: var(--muted);
          font-size: 13px;
        }
        .toc-list {
          display: grid;
          gap: 8px;
          overflow: hidden;
        }
        .toc-link,
        .toc-empty {
          color: var(--muted);
          font-size: 13px;
          line-height: 1.45;
          transform-origin: left center;
          transition: color 260ms var(--ease), transform 320ms var(--ease);
        }
        .toc-link:hover {
          color: var(--ink);
          transform: scale(1.045) translateX(3px);
        }
        .toc-list:has(.toc-link:hover) .toc-link:not(:hover) {
          transform: translateY(1px);
        }
        .toc-level-3,
        .toc-level-4 {
          padding-left: 14px;
        }
        .article-content {
          max-width: 720px;
          padding-left: 24px;
        }
        .article-header {
          padding-bottom: 0;
        }
        .article-header h1 {
          font-size: 42px;
          line-height: 1.14;
          margin-bottom: 0.22em;
        }
        .article-source-link {
          margin: 18px 0 0;
          color: var(--muted);
          font-size: 16px;
          font-weight: 400;
          line-height: 1.9;
        }
        .article-keyline {
          margin: 28px 0 0;
          color: var(--ink);
          font-size: 16px;
          font-weight: 400;
          line-height: 1.9;
        }
        .article-divider {
          margin: 28px 0 0;
          border-top: 1px solid var(--line);
        }
        .article-source-link a {
          border-bottom: 1px solid currentColor;
        }
        .article-source-link a:hover {
          color: var(--ink);
        }
        .article-source-link + .body {
          margin-top: 24px;
        }
        .body { margin-top: 44px; }
        .body h2,
        .body h3,
        .body h4 {
          scroll-margin-top: 92px;
          margin: 2.2em 0 0.72em;
          color: var(--ink);
          font-weight: 300;
          line-height: 1.28;
          letter-spacing: 0;
        }
        .body h2 { font-size: 24px; padding-top: 0.25em; }
        .body h3 { font-size: 20px; }
        .body h4 { font-size: 18px; }
        .body p {
          margin: 1.12em 0;
          color: var(--ink);
          font-size: 16px;
          font-weight: 400;
          line-height: 1.9;
          text-indent: 0;
        }
        .contact-grid {
          display: grid;
          grid-template-columns: repeat(2, minmax(0, 1fr));
          gap: 18px;
          padding-top: 42px;
        }
        .contact-block {
          min-height: 150px;
          padding: 22px;
          border: 1px solid var(--line);
          border-radius: var(--radius);
          background: var(--panel);
        }
        .contact-block span {
          color: var(--muted);
          font-size: 13px;
        }
        .contact-block p {
          margin: 18px 0 0;
          color: var(--ink);
          font-size: 20px;
          font-weight: 300;
        }
        .contact-block p > a {
          border-bottom: 1px solid currentColor;
        }
        .social-list {
          display: grid;
          gap: 12px;
          margin-top: 20px;
        }
        .social-item {
          display: grid;
          grid-template-columns: 34px minmax(0, 1fr);
          gap: 12px;
          align-items: center;
          width: 100%;
          border: 1px solid var(--line);
          border-radius: 12px;
          padding: 10px 12px;
          background: transparent;
          color: var(--ink);
          cursor: pointer;
          text-align: left;
          transition: border-color 180ms var(--ease), background 180ms var(--ease), transform 180ms var(--ease);
        }
        .social-item:hover {
          border-color: var(--ink);
          background: var(--paper);
          transform: translateY(-1px);
        }
        .social-item img {
          width: 34px;
          height: 34px;
          border-radius: 9px;
          object-fit: cover;
        }
        .social-item span {
          color: var(--ink);
          font-size: 14px;
          line-height: 1.45;
        }
        .copy-social[data-copied="true"] span::after {
          content: " 已复制";
          color: var(--muted);
        }
        .download-link {
          display: inline-block;
          border-bottom: 1px solid currentColor;
        }
        .download-link.is-disabled {
          color: var(--muted);
          border-bottom: 0;
        }
        .empty { color: var(--muted); margin: 0.5em 0; }
        @media (max-width: 900px) {
          .about-hero,
          .feature-primary,
          .article-read,
          .archive-controls {
            grid-template-columns: 1fr;
          }
          .about-portrait { justify-items: start; }
          .selected-layout,
          .archive-grid-view.is-active,
          .field-grid {
            grid-template-columns: repeat(2, minmax(0, 1fr));
          }
          .article-visual {
            position: relative;
            top: auto;
            max-height: none;
            overflow: visible;
          }
          .article-content {
            max-width: none;
            padding-left: 0;
          }
        }
        @media (max-width: 640px) {
          .site-header {
            position: relative;
            align-items: flex-start;
            flex-direction: column;
            gap: 12px;
            width: min(100% - 30px, 1120px);
            padding: 18px 0 16px;
          }
          .main-nav {
            width: 100%;
            justify-content: flex-start;
            gap: 14px;
            overflow-x: auto;
            padding-bottom: 2px;
          }
          .shell { width: min(100% - 30px, 1120px); padding-top: 28px; padding-bottom: 72px; }
          .site-footer {
            width: min(100% - 30px, 1120px);
            flex-direction: column;
            align-items: flex-start;
            gap: 6px;
          }
          .about-hero, .page-title { padding: 42px 0 32px; }
          .about-hero h1, .page-title h1, .article-header h1 { font-size: 34px; line-height: 1.16; }
          .hero-bio { font-size: 16px; line-height: 1.78; }
          .about-dashboard,
          .selected-layout,
          .archive-grid-view.is-active,
          .field-grid,
          .contact-grid {
            grid-template-columns: 1fr;
          }
          .about-dashboard div,
          .about-dashboard div:last-child {
            border-right: 0;
            border-bottom: 1px solid var(--line);
            padding: 20px 0;
          }
          .about-dashboard div:last-child { border-bottom: 0; }
          .feature-card h2 { font-size: 24px; }
          .feature-primary h2 { font-size: 30px; }
          .archive-item { grid-template-columns: 92px minmax(0, 1fr); gap: 14px; padding: 22px 0; }
          .archive-image { grid-row: 1; }
          .archive-main h3 { font-size: 19px; }
          .view-toggle { justify-content: flex-start; }
          .body h2 { font-size: 22px; }
          .body h3 { font-size: 19px; }
          .body p { font-size: 16px; line-height: 1.85; }
        }
        @media (prefers-reduced-motion: reduce) {
          html { scroll-behavior: auto; }
          *,
          *::before,
          *::after {
            transition-duration: 1ms !important;
            animation-duration: 1ms !important;
          }
        }
        """,
        encoding="utf-8",
    )


def keywords_from_query(query: str) -> list[str]:
    cleaned = re.sub(r"(岗位|申请|一个|一份|类|撰稿人|作品集|请|帮我|整理|相关|文档)", " ", query)
    parts = split_tags(cleaned)
    more = re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z0-9]{2,}", cleaned)
    result = []
    for item in [*parts, *more]:
        item = item.strip()
        if item and item not in result:
            result.append(item)
    result = expand_keywords(result)
    return result or [query.strip()]


def expand_keywords(keywords: list[str]) -> list[str]:
    groups = {
        "文旅": ["文旅", "文化", "旅游", "旅行", "城市", "生活方式", "Travel", "Lifestyle"],
        "人物": ["人物", "访谈", "专访", "Profiles", "Interviews"],
        "品牌": ["品牌", "商业", "Brand", "Features"],
        "建筑": ["建筑", "设计", "空间", "Design", "Architecture"],
    }
    expanded = []
    for keyword in keywords:
        additions = []
        for trigger, values in groups.items():
            if trigger in keyword:
                additions.extend(values)
        additions.append(keyword)
        for item in additions:
            if item and item not in expanded:
                expanded.append(item)
    return expanded


def match_articles(articles: list[dict], query: str) -> list[dict]:
    if is_all_portfolio_query(query):
        return sort_articles_by_type(articles)
    keywords = keywords_from_query(query)
    matched = []
    for article in articles:
        haystack = " ".join(
            [
                str(article.get("title", "")),
                str(article.get("writing_type", "")),
                " ".join(article.get("tags", [])),
                str(article.get("media", "")),
                str(article.get("notes", "")),
                first_paragraph(article.get("body", ""), 500),
            ]
        )
        score = sum(1 for keyword in keywords if keyword and keyword in haystack)
        if score:
            item = dict(article)
            item["_score"] = score
            matched.append(item)
    matched.sort(key=lambda item: (item["_score"], str(item.get("publish_date") or "")), reverse=True)
    return matched


def is_all_portfolio_query(query: str) -> bool:
    return bool(re.search(r"(全部|全量|完整|所有)", normalize_spaces(query)))


def wants_grouped_pdf(query: str) -> bool:
    return is_all_portfolio_query(query) or "分门别类" in normalize_spaces(query)


def sort_articles_by_type(articles: list[dict]) -> list[dict]:
    return sorted(
        (dict(article) for article in articles),
        key=lambda item: (
            type_sort_index(item.get("writing_type")),
            date_sort_desc(item.get("publish_date")),
        ),
    )


def article_group_name(article: dict) -> str:
    return canonical_writing_type(article.get("writing_type") or "未分类")


def make_summary(article: dict, query: str) -> str:
    tags = "、".join(article.get("tags", [])) or "未标注"
    writing_type = display_writing_type(article.get("writing_type") or "未分类")
    media = article.get("media") or "未标注媒体"
    return f"这篇作品属于{writing_type}类撰稿，发布于{media}，主题标签为{tags}。用于“{query}”时，可重点呈现其选题判断、信息整理和文本表达能力。"


def format_hash_tags(article: dict) -> str:
    tags = []
    for tag in article.get("tags", []):
        clean = normalize_spaces(tag).lstrip("#")
        if clean:
            tags.append(f"#{clean}")
    if not tags:
        writing_type = display_writing_type(article.get("writing_type") or "")
        if writing_type:
            tags.append(f"#{writing_type}")
    return " ".join(tags) or "#未标注"


def display_writing_type(value: object) -> str:
    text = normalize_spaces(value)
    text = re.sub(r"（[^）]*[A-Za-z][^）]*）", "", text)
    text = re.sub(r"\([^)]*[A-Za-z][^)]*\)", "", text)
    return text.strip() or "未分类"


def make_excerpt(article: dict, query: str, limit: int = 420) -> str:
    body = article.get("body", "")
    keywords = keywords_from_query(query)
    paragraphs = [
        p.strip()
        for p in body.split("\n\n")
        if p.strip() and not re.match(r"^\s{0,3}#{1,6}\s+", p.strip())
    ]
    for keyword in keywords:
        for paragraph in paragraphs:
            if keyword in paragraph:
                return trim_text(paragraph, limit)
    return trim_text(" ".join(paragraphs[:3]), limit)


class PortfolioDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, *, body_start_page: int = 2, **kwargs) -> None:
        super().__init__(filename, **kwargs)
        self.body_start_page = body_start_page
        self.article_pages: dict[str, int] = {}

    def afterFlowable(self, flowable) -> None:
        anchor = getattr(flowable, "_bookmark_name", None)
        title = getattr(flowable, "_bookmark_title", "")
        if not anchor:
            return
        self.canv.bookmarkPage(anchor)
        self.article_pages[anchor] = max(1, self.page - self.body_start_page + 1)
        if title:
            self.canv.addOutlineEntry(title, anchor, level=0, closed=False)


def portfolio_subject(query: str) -> str:
    if is_all_portfolio_query(query):
        return "全部"
    text = normalize_spaces(query)
    text = re.sub(r"(作品集|撰稿人|岗位|申请|一个|一份|相关|文档|请|帮我|整理)", "", text)
    text = text.replace("类", "")
    text = sanitize_filename_part(text)
    return text or "综合"


def portfolio_pdf_title(subject: str, all_export: bool, full_text: bool = False) -> str:
    suffix = "全文版" if full_text else "摘要版"
    if all_export:
        return f"Su FAN 全部作品集（{suffix}）"
    return f"Su FAN {subject}类作品集（{suffix}）"


def portfolio_pdf_filename(subject: str, all_export: bool, full_text: bool = False) -> str:
    date_prefix = dt.date.today().strftime("%Y%m%d")
    suffix = "全文版" if full_text else "摘要版"
    if all_export:
        return f"{date_prefix}-SuFAN全部作品集-{suffix}.pdf"
    return f"{date_prefix}-SuFAN{subject}类作品集-{suffix}.pdf"


def roman_number(number: int) -> str:
    values = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]
    result = []
    for value, symbol in values:
        while number >= value:
            result.append(symbol)
            number -= value
    return "".join(result)


def portfolio_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont(PDF_FONT, 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    page = canvas.getPageNumber()
    if page < doc.body_start_page:
        label = roman_number(page)
    else:
        label = str(page - doc.body_start_page + 1)
    canvas.drawRightString(A4[0] - 18 * mm, 10 * mm, label)
    canvas.restoreState()


def article_anchor(index: int) -> str:
    return f"article_{index}"


def markdown_heading(block: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+)$", block.strip())
    if not match or "\n" in block:
        return None
    return len(match.group(1)), match.group(2).strip()


def wants_full_text_pdf(query: str) -> bool:
    return "全文版" in normalize_spaces(query)


def make_pdf(query: str, full_text: bool = False) -> Path:
    ensure_dirs()
    full_text = full_text or wants_full_text_pdf(query)
    articles = load_articles(public_only=False) if full_text else load_summary_articles(public_only=False)
    matched = match_articles(articles, query)
    if not matched:
        raise SystemExit(f"没有找到与“{query}”匹配的已入库作品。请检查撰稿类型或主题标签。")

    regular_font, bold_font = register_pdf_fonts()
    all_export = is_all_portfolio_query(query)
    grouped = wants_grouped_pdf(query)
    subject = portfolio_subject(query)
    output = PDF_DIR / portfolio_pdf_filename(subject, all_export, full_text)
    pdf_title = portfolio_pdf_title(subject, all_export, full_text)

    def create_doc(path: Path) -> PortfolioDocTemplate:
        frame = Frame(18 * mm, 18 * mm, A4[0] - 36 * mm, A4[1] - 36 * mm, id="main")
        doc = PortfolioDocTemplate(
            str(path),
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=pdf_title,
            body_start_page=2,
        )
        doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=portfolio_footer)])
        return doc

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "ChineseBase",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10.5,
        leading=17,
        textColor=colors.HexColor("#172026"),
        spaceAfter=8,
    )
    title_style = ParagraphStyle(
        "TitleCN",
        parent=base,
        fontName=bold_font,
        fontSize=24,
        leading=30,
        spaceAfter=12,
        textColor=colors.HexColor("#111827"),
    )
    cover_title = ParagraphStyle(
        "CoverTitle",
        parent=title_style,
        fontSize=26,
        leading=32,
        spaceAfter=12,
    )
    h2 = ParagraphStyle(
        "ArticleTitle",
        parent=base,
        fontName=bold_font,
        fontSize=16,
        leading=22,
        spaceBefore=0,
        spaceAfter=8,
        textColor=colors.HexColor("#8A4B38"),
    )
    small = ParagraphStyle(
        "SmallCN",
        parent=base,
        fontSize=9,
        leading=14,
        textColor=colors.HexColor("#536068"),
    )
    source_link_style = ParagraphStyle(
        "SourceLink",
        parent=small,
        fontSize=9,
        leading=13,
        spaceAfter=6,
        textColor=colors.HexColor("#536068"),
    )
    toc_style = ParagraphStyle(
        "TocLine",
        parent=base,
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#172026"),
    )
    toc_page_style = ParagraphStyle(
        "TocPage",
        parent=toc_style,
        alignment=2,
    )
    toc_group_style = ParagraphStyle(
        "TocGroup",
        parent=toc_style,
        fontName=bold_font,
        fontSize=10.5,
        leading=16,
        spaceBefore=8,
        textColor=colors.HexColor("#8A4B38"),
    )
    section_style = ParagraphStyle(
        "SectionLabel",
        parent=base,
        fontName=bold_font,
        fontSize=11,
        leading=16,
        spaceBefore=8,
        textColor=colors.HexColor("#111827"),
    )
    body_heading_2 = ParagraphStyle(
        "BodyHeading2",
        parent=base,
        fontName=bold_font,
        fontSize=12,
        leading=18,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor("#111827"),
    )
    body_heading_3 = ParagraphStyle(
        "BodyHeading3",
        parent=body_heading_2,
        fontSize=10.5,
        leading=17,
        spaceBefore=10,
        spaceAfter=5,
    )

    def build_story(toc_pages: dict[str, int] | None = None) -> list:
        toc_pages = toc_pages or {}
        media_list = "、".join(media_names(load_summary_articles(public_only=True))) or "暂无"
        profile = [
            "现为自由撰稿人，base 北京。",
            "中英文流利，能从容应对双语采访与撰稿任务。",
            "教育背景：经济学博士，北京大学（2025）；联合培养博士研究生，巴黎第一大学（2023）；经济学、社会学学士，北京大学（2020）。",
        ]
        story = [
            Paragraph(html_escape(pdf_title), cover_title),
        ]
        for line in profile:
            story.append(Paragraph(html_escape(line), base))
        story.extend(
            [
                Paragraph(html_escape(f"合作媒体：{media_list}"), small),
                Spacer(1, 10),
                Paragraph("目录", title_style),
            ]
        )
        rows = []
        toc_commands = [
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#D8D4CA")),
        ]
        current_group = ""
        for index, article in enumerate(matched, start=1):
            group_name = article_group_name(article)
            if grouped and group_name != current_group:
                row_index = len(rows)
                rows.append([Paragraph(html_escape(group_name), toc_group_style), Paragraph("", toc_page_style)])
                toc_commands.append(("SPAN", (0, row_index), (1, row_index)))
                toc_commands.append(("LINEBELOW", (0, row_index), (-1, row_index), 0.5, colors.HexColor("#111827")))
                current_group = group_name
            anchor = article_anchor(index)
            label = f"{article.get('media') or ''}｜{article.get('title') or ''}"
            page_label = str(toc_pages.get(anchor, ""))
            rows.append(
                [
                    Paragraph(f'<a href="#{anchor}">{html_escape(label)}</a>', toc_style),
                    Paragraph(f'<a href="#{anchor}">{html_escape(page_label)}</a>', toc_page_style),
                ]
            )
        toc = Table(rows, colWidths=[140 * mm, 20 * mm])
        toc.setStyle(TableStyle(toc_commands))
        story.extend([toc, PageBreak()])

        current_group = ""
        for index, article in enumerate(matched, start=1):
            if index > 1:
                story.append(PageBreak())
            group_name = article_group_name(article)
            if grouped and group_name != current_group:
                story.append(Paragraph(html_escape(group_name), section_style))
                current_group = group_name
            anchor = article_anchor(index)
            title = str(article.get("title") or "")
            title_para = Paragraph(html_escape(title), h2)
            title_para._bookmark_name = anchor
            title_para._bookmark_title = title
            story.append(title_para)
            meta_parts = [
                normalize_spaces(article.get("media")),
                normalize_spaces(article.get("publish_date")),
                display_writing_type(article.get("writing_type") or ""),
                *[normalize_spaces(tag) for tag in article.get("tags", []) if normalize_spaces(tag)],
            ]
            meta = " | ".join(part for part in meta_parts if part)
            story.append(Paragraph(html_escape(meta), small))
            if not full_text and confirmed_summary(article):
                keyline = normalize_spaces(article.get("keyline"))
                if keyline:
                    story.append(Paragraph(html_escape(keyline), base))
            url = original_url(article)
            if url:
                source_text = (
                    f'完整内容请点击 <a href="{html_escape(url)}">阅读原文</a>，以下为精选节选'
                    if not full_text
                    else f'<a href="{html_escape(url)}">阅读原文</a>'
                )
                story.append(
                    Paragraph(
                        source_text,
                        source_link_style,
                    )
                )
            body_source = str(article.get("body") or "")
            if not full_text and not confirmed_summary(article):
                body_source = "摘要待确认。请点击阅读原文查看完整作品。"
            for paragraph in [p.strip() for p in body_source.split("\n\n") if p.strip()]:
                lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
                first_line = lines[0] if lines else ""
                heading = markdown_heading(first_line)
                if heading:
                    level, heading_text = heading
                    story.append(Paragraph(html_escape(heading_text), body_heading_2 if level <= 2 else body_heading_3))
                    rest = "\n".join(lines[1:]).strip()
                    if rest:
                        story.append(Paragraph(html_escape(rest), base))
                    continue
                story.append(Paragraph(html_escape(paragraph), base))
        return story

    TMP.mkdir(exist_ok=True)
    draft = TMP / "portfolio_pdf_draft.pdf"
    draft_doc = create_doc(draft)
    draft_doc.build(build_story())
    final_doc = create_doc(output)
    final_doc.build(build_story(draft_doc.article_pages))

    print(f"PDF 已生成：{output}")
    return output


def register_pdf_fonts() -> tuple[str, str]:
    candidates = [
        (
            Path("/System/Library/Fonts/STHeiti Light.ttc"),
            Path("/System/Library/Fonts/STHeiti Medium.ttc"),
        ),
        (
            Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
            Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
        ),
    ]
    for regular, bold in candidates:
        if not regular.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(PDF_FONT, str(regular)))
            if bold.exists():
                pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(bold)))
            else:
                pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(regular)))
            return PDF_FONT, PDF_FONT_BOLD
        except Exception:
            continue
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light", "STSong-Light"


def run_all(query: str | None) -> None:
    ingest()
    build_site()
    if query:
        make_pdf(query)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="作品集整理流程")
    sub = parser.add_subparsers(dest="command", required=True)

    init_cmd = sub.add_parser("init", help="初始化台账模板和文件夹")
    init_cmd.add_argument("--force", action="store_true", help="覆盖已有台账模板")

    sub.add_parser("ingest", help="整理入库")
    sub.add_parser("site", help="更新静态网站")
    sub.add_parser("summaries", help="生成缺失的摘要草稿并列出待审核作品")

    pdf_cmd = sub.add_parser("pdf", help="生成 PDF 作品集")
    pdf_cmd.add_argument("query", help="岗位或主题，例如：文旅类撰稿人岗位")

    pdf_full_cmd = sub.add_parser("pdf-full", help="生成全文版 PDF 作品集")
    pdf_full_cmd.add_argument("query", help="岗位或主题，例如：全部作品集，分门别类")

    all_cmd = sub.add_parser("all", help="入库、更新网站，并可生成 PDF")
    all_cmd.add_argument("query", nargs="?", help="岗位或主题，例如：文旅类撰稿人岗位")

    publish_cmd = sub.add_parser("publish", help="更新静态网站、提交并推送到 GitHub 备份")
    publish_cmd.add_argument("message", help="本次备份说明，例如：新增某某作品")

    sub.add_parser("deploy-vercel", help="更新静态网站并通过 Vercel CLI 发布全球版")
    sub.add_parser("deploy-tencent-server", help="更新静态网站并部署到腾讯云服务器")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.command == "init":
        init_project(force=args.force)
    elif args.command == "ingest":
        ingest()
    elif args.command == "site":
        build_site()
    elif args.command == "summaries":
        summaries_status()
    elif args.command == "pdf":
        make_pdf(args.query)
    elif args.command == "pdf-full":
        make_pdf(args.query, full_text=True)
    elif args.command == "all":
        run_all(args.query)
    elif args.command == "publish":
        publish_site(args.message)
    elif args.command == "deploy-vercel":
        deploy_vercel()
    elif args.command == "deploy-tencent-server":
        deploy_tencent_server()


if __name__ == "__main__":
    main()
